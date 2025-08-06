"""
Unified Template Processor for Multi-Item Quote Generation

This processor handles both single-item and multi-item quotes using a single master template
and model-specific configuration files.
"""

import os
import json
import re
from pathlib import Path
from typing import Dict, Any, List, Optional, TYPE_CHECKING
from datetime import datetime
import logging

if TYPE_CHECKING:
    from docx import Document

try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)

class UnifiedTemplateProcessor:
    """
    Processes the master template with dynamic content based on model configurations.
    Supports both single-item and multi-item quotes.
    """
    
    def __init__(self, templates_dir: Optional[str] = None):
        """
        Initialize the unified template processor.
        
        Args:
            templates_dir: Path to the unified templates directory
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required. Install with: pip install python-docx")
            
        if templates_dir is None:
            templates_dir = os.path.join(os.path.dirname(__file__))
        
        self.templates_dir = Path(templates_dir)
        self.configs_dir = self.templates_dir / 'configs'
        self.master_template_path = self.templates_dir / 'master_template.docx'
        
        # Load template schema
        self.schema = self._load_schema()
    
    def _load_schema(self) -> Dict[str, Any]:
        """Load the template schema configuration."""
        schema_path = self.configs_dir / 'template_schema.json'
        try:
            with open(schema_path, 'r') as f:
                return json.load(f)
        except Exception as e:
            logger.warning(f"Could not load schema: {e}. Using defaults.")
            return {"multi_item_support": {"enabled": True}}
    
    def _load_model_config(self, model: str) -> Dict[str, Any]:
        """
        Load configuration for a specific model.
        
        Args:
            model: Model name (e.g., 'LS2000', 'LS7000')
            
        Returns:
            Model configuration dictionary
        """
        config_path = self.configs_dir / f'{model}_config.json'
        try:
            with open(config_path, 'r') as f:
                return json.load(f)
        except Exception as e:
            logger.error(f"Could not load config for model {model}: {e}")
            return self._get_default_config(model)
    
    def _get_default_config(self, model: str) -> Dict[str, Any]:
        """Get default configuration when model config is not found."""
        return {
            "model": model,
            "description": f"{model} Level Switch Quote",
            "technical_specifications": [
                {"label": "Supply Voltage", "value": "{{supply_voltage}}", "static": False},
                {"label": "Output", "value": "10 Amp SPDT Relay", "static": True},
                {"label": "Process Connection", "value": "{{pc_size}} {{pc_type}}", "static": False},
                {"label": "Probe", "value": "{{probe_size}}\" Diameter {{probe_material}} x {{probe_length}}\"", "static": False}
            ],
            "optional_sections": {"notes": {"enabled": False, "content": []}},
            "default_values": {}
        }
    
    def _build_technical_specifications(self, config: Dict[str, Any], variables: Dict[str, str]) -> str:
        """
        Build the technical specifications list from config and variables.
        
        Args:
            config: Model configuration
            variables: Template variables
            
        Returns:
            Formatted specifications as string
        """
        specs = config.get('technical_specifications', [])
        spec_lines = []
        
        for spec in specs:
            label = spec.get('label', '')
            value_template = spec.get('value', '')
            
            # Replace variables in the value template
            value = value_template
            for var_name, var_value in variables.items():
                placeholder = f"{{{{{var_name}}}}}"
                value = value.replace(placeholder, str(var_value))
            
            spec_lines.append(f"• {label}: {value}")
        
        return '\n'.join(spec_lines)
    
    def _build_technical_specifications_bullets(self, config: Dict[str, Any], variables: Dict[str, str]) -> str:
        """
        Build technical specifications as bullet points matching original template format.
        
        Args:
            config: Model configuration  
            variables: Template variables
            
        Returns:
            Formatted bullet-point specifications as string
        """
        specs = config.get('technical_specifications', [])
        spec_lines = []
        
        for spec in specs:
            label = spec.get('label', '')
            value_template = spec.get('value', '')
            
            # Replace variables in the value template
            value = value_template
            for var_name, var_value in variables.items():
                placeholder = f"{{{{{var_name}}}}}"
                if var_value:  # Only replace if value is not empty
                    value = value.replace(placeholder, str(var_value))
            
            # Clean up any remaining placeholders
            import re
            value = re.sub(r'\{\{[^}]+\}\}', '', value)
            
            # Only add non-empty specifications
            if value.strip():
                spec_lines.append(f"• {label}: {value}")
        
        return '\n'.join(spec_lines)
    
    def _build_items_section(self, quote_items: List[Dict[str, Any]], variables: Dict[str, str]) -> str:
        """
        Build the items section for single or multi-item quotes.
        
        Args:
            quote_items: List of quote items
            variables: Template variables
            
        Returns:
            Formatted items section as string
        """
        if len(quote_items) == 1:
            return self._build_single_item_section(quote_items[0], variables)
        else:
            return self._build_multi_item_section(quote_items, variables)
    
    def _build_single_item_section(self, item: Dict[str, Any], variables: Dict[str, str]) -> str:
        """Build section for single item quote in original template format."""
        part_number = item.get('part_number', '')
        quantity = item.get('quantity', 1)
        
        # Get model from part number
        model = self._extract_model_from_part_number(part_number)
        config = self._load_model_config(model)
        
        # Merge item data into variables
        item_variables = variables.copy()
        if 'data' in item:
            item_variables.update(item['data'])
        item_variables.update({
            'part_number': part_number,
            'quantity': str(quantity),
            'model': model
        })
        
        # Build single item format to match original template exactly
        unit_price = item_variables.get('unit_price', '0.00')
        if not unit_price.startswith('$'):
            unit_price = f"${unit_price}"
        
        section = f"{quantity} QTY {part_number}             {unit_price}    EACH \n\n"
        
        # Build technical specifications as bullet points to match original format
        tech_specs = self._build_technical_specifications_bullets(config, item_variables)
        section += tech_specs
        
        return section
    
    def _build_multi_item_section(self, quote_items: List[Dict[str, Any]], variables: Dict[str, str]) -> str:
        """Build section for multi-item quote with simplified format."""
        sections = []
        
        for i, item in enumerate(quote_items, 1):
            part_number = item.get('part_number', '')
            quantity = item.get('quantity', 1)
            
            # Get model from part number
            model = self._extract_model_from_part_number(part_number)
            config = self._load_model_config(model)
            
            # Merge item data into variables
            item_variables = variables.copy()
            if 'data' in item:
                item_variables.update(item['data'])
            item_variables.update({
                'part_number': part_number,
                'quantity': str(quantity),
                'model': model
            })
            
            # Get unit price
            if item.get('type') == 'main':
                unit_price = item.get('data', {}).get('total_price', 0.0)
            else:
                unit_price = item.get('data', {}).get('pricing', {}).get('total_price', 0.0)
            
            # Build item section in clean format
            item_section = f"Item {i}: {quantity} QTY {part_number}             ${unit_price:.2f} EACH\n"
            
            # Build technical specifications for this item
            tech_specs = self._build_technical_specifications_bullets(config, item_variables)
            if tech_specs:
                item_section += tech_specs + "\n"
            
            sections.append(item_section)
        
        return "\n\n".join(sections)
    
    def _build_quote_summary_table(self, quote_items: List[Dict[str, Any]]) -> str:
        """Build quote summary table for multi-item quotes."""
        if len(quote_items) <= 1:
            return ""  # No summary needed for single item
        
        # Simple text-based summary
        grand_total = 0.0
        
        for item in quote_items:
            quantity = item.get('quantity', 1)
            
            # Get unit price
            if item.get('type') == 'main':
                unit_price = item.get('data', {}).get('total_price', 0.0)
            else:
                unit_price = item.get('data', {}).get('pricing', {}).get('total_price', 0.0)
            
            total_price = unit_price * quantity
            grand_total += total_price
        
        summary = f"\nQuote Total: ${grand_total:.2f}"
        return summary
    
    def _extract_model_from_part_number(self, part_number: str) -> str:
        """Extract model from part number (e.g., 'LS2000-115VAC-S-12' -> 'LS2000')."""
        if '-' in part_number:
            return part_number.split('-')[0]
        elif len(part_number) >= 6:
            return part_number[:6]
        else:
            return part_number
    
    def _build_optional_notes_section(self, quote_items: List[Dict[str, Any]]) -> str:
        """Build optional notes section based on item models."""
        notes_sections = []
        processed_models = set()
        
        for item in quote_items:
            part_number = item.get('part_number', '')
            model = self._extract_model_from_part_number(part_number)
            
            if model in processed_models:
                continue
            processed_models.add(model)
            
            config = self._load_model_config(model)
            notes_config = config.get('optional_sections', {}).get('notes', {})
            
            if notes_config.get('enabled', False):
                notes_content = notes_config.get('content', [])
                if notes_content:
                    if not notes_sections:  # First notes section
                        notes_sections.append("Notes:")
                    for note in notes_content:
                        notes_sections.append(f"• {note}")
        
        return "\n".join(notes_sections) if notes_sections else ""
    
    def process_unified_template(self, quote_items: List[Dict[str, Any]], variables: Dict[str, Any]) -> Optional[Any]:
        """
        Process the master template with quote items and variables.
        
        Args:
            quote_items: List of quote items with 'part_number', 'quantity', 'data'
            variables: Base template variables (customer, date, etc.)
            
        Returns:
            Processed Document object, or None if failed
        """
        if not self.master_template_path.exists():
            logger.error(f"Master template not found: {self.master_template_path}")
            return None
        
        try:
            logger.info(f"Processing unified template: {self.master_template_path}")
            
            # Load the master template
            doc = Document(str(self.master_template_path))
            
            # Convert variables to strings
            str_variables = {k: str(v) if v is not None else "" for k, v in variables.items()}
            
            # Determine if this is a multi-item quote
            is_multi_item = len(quote_items) > 1
            str_variables['has_multiple_items'] = str(is_multi_item).lower()
            str_variables['item_count'] = str(len(quote_items))
            
            # Build dynamic content sections
            if is_multi_item:
                str_variables['model_description'] = "Multi-Item Quote"
            else:
                # Get model description from first item
                model = self._extract_model_from_part_number(quote_items[0].get('part_number', ''))
                config = self._load_model_config(model)
                str_variables['model_description'] = config.get('description', f"{model} Level Switch Quote")
            
            # Build items section
            str_variables['items_section'] = self._build_items_section(quote_items, str_variables)
            
            # Build quote summary table
            str_variables['quote_summary_table'] = self._build_quote_summary_table(quote_items)
            
            # Build optional notes section
            str_variables['optional_notes_section'] = self._build_optional_notes_section(quote_items)
            
            # Process conditional content
            self._process_conditional_content_document(doc, str_variables)
            
            # Process all paragraphs for variable replacement
            for paragraph in doc.paragraphs:
                self._replace_variables_in_paragraph(paragraph, str_variables)
            
            # Process all tables
            for table in doc.tables:
                self._replace_variables_in_table(table, str_variables)
            
            # Process headers and footers
            for section in doc.sections:
                header = section.header
                for paragraph in header.paragraphs:
                    self._replace_variables_in_paragraph(paragraph, str_variables)
                
                footer = section.footer
                for paragraph in footer.paragraphs:
                    self._replace_variables_in_paragraph(paragraph, str_variables)
            
            logger.info("Unified template processing completed successfully")
            return doc
            
        except Exception as e:
            logger.error(f"Error processing unified template: {e}")
            import traceback
            logger.error(f"Traceback: {traceback.format_exc()}")
            return None
    
    def _replace_variables_in_paragraph(self, paragraph, variables: Dict[str, str]) -> None:
        """Replace template variables in a paragraph while preserving formatting."""
        full_text = paragraph.text
        variable_pattern = r'\{\{([^}]+)\}\}'
        matches = re.findall(variable_pattern, full_text)
        
        if not matches:
            return
        
        new_text = full_text
        for var_name in matches:
            placeholder = f"{{{{{var_name}}}}}"
            value = variables.get(var_name, f"{{{{MISSING: {var_name}}}}}")
            new_text = new_text.replace(placeholder, str(value))
        
        paragraph.clear()
        paragraph.add_run(new_text)
    
    def _replace_variables_in_table(self, table, variables: Dict[str, str]) -> None:
        """Replace template variables in a table while preserving formatting."""
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    self._replace_variables_in_paragraph(paragraph, variables)
    
    def _process_conditional_content_document(self, doc, variables: Dict[str, str]) -> None:
        """Process conditional content markers in the document."""
        # Process paragraphs
        for paragraph in doc.paragraphs:
            self._process_conditional_content_paragraph(paragraph, variables)
        
        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._process_conditional_content_paragraph(paragraph, variables)
    
    def _process_conditional_content_paragraph(self, paragraph, variables: Dict[str, str]) -> None:
        """Process conditional content in a single paragraph."""
        full_text = paragraph.text
        
        # Process single item conditionals
        single_item_pattern = r'\{\{if_single_item:([^}]+)\}\}'
        single_item_matches = list(re.finditer(single_item_pattern, full_text))
        
        # Process multi-item conditionals
        multi_item_pattern = r'\{\{if_multiple_items:([^}]+)\}\}'
        multi_item_matches = list(re.finditer(multi_item_pattern, full_text))
        
        # Combine and sort matches
        all_matches = []
        for match in single_item_matches:
            all_matches.append(('single_item', match))
        for match in multi_item_matches:
            all_matches.append(('multi_item', match))
        
        # Sort by position in reverse order
        all_matches.sort(key=lambda x: x[1].start(), reverse=True)
        
        if not all_matches:
            return
        
        # Process conditionals
        is_multi_item = variables.get('has_multiple_items', 'false').lower() == 'true'
        
        for match_type, match in all_matches:
            start_pos = match.start()
            end_pos = match.end()
            content = match.group(1)
            
            if match_type == 'single_item':
                replacement = content if not is_multi_item else ""
            elif match_type == 'multi_item':
                replacement = content if is_multi_item else ""
            else:
                replacement = ""
            
            # Replace the conditional block
            paragraph.clear()
            new_text = full_text[:start_pos] + replacement + full_text[end_pos:]
            paragraph.add_run(new_text)
            full_text = new_text
    
    def save_document(self, doc: Any, output_path: str) -> bool:
        """Save the processed document."""
        try:
            doc.save(output_path)
            return True
        except Exception as e:
            logger.error(f"Error saving document to {output_path}: {e}")
            return False


def generate_unified_quote(
    quote_items: List[Dict[str, Any]],
    customer_name: str,
    attention_name: str,
    quote_number: str,
    output_path: str,
    employee_info: Optional[Dict[str, str]] = None,
    **kwargs
) -> bool:
    """
    Generate a quote using the unified template system.
    
    Args:
        quote_items: List of quote items
        customer_name: Customer company name
        attention_name: Contact person name
        quote_number: Quote number
        output_path: Output file path
        employee_info: Employee information
        **kwargs: Additional variables
        
    Returns:
        True if successful, False otherwise
    """
    if not DOCX_AVAILABLE:
        logger.error("python-docx not available")
        return False
    
    try:
        # Create processor
        processor = UnifiedTemplateProcessor()
        
        # Get employee information
        employee_name = kwargs.get('employee_name', 'John Nicholosi')
        employee_phone = kwargs.get('employee_phone', '(713) 467-4438')
        employee_email = kwargs.get('employee_email', 'John@babbitt.us')
        
        if employee_info:
            employee_name = employee_info.get('name', employee_name)
            employee_phone = employee_info.get('phone', employee_phone)
            employee_email = employee_info.get('email', employee_email)
        
        # Prepare base variables
        variables = {
            'date': datetime.now().strftime("%B %d, %Y"),
            'customer_name': customer_name,
            'attention_name': attention_name,
            'quote_number': quote_number,
            'employee_name': employee_name,
            'employee_phone': employee_phone,
            'employee_email': employee_email,
            'lead_time': kwargs.get('lead_time', 'In Stock')
        }
        
        # Add any additional variables
        variables.update(kwargs)
        
        # Process the template
        doc = processor.process_unified_template(quote_items, variables)
        if not doc:
            logger.error("Failed to process unified template")
            return False
        
        # Save the document
        success = processor.save_document(doc, output_path)
        if success:
            logger.info(f"Unified quote generated successfully: {output_path}")
        
        return success
        
    except Exception as e:
        logger.error(f"Error generating unified quote: {e}")
        return False


# Test function
def test_unified_processing():
    """Test the unified template processing with sample data."""
    if not DOCX_AVAILABLE:
        print("❌ python-docx not available")
        return False
    
    # Test with single item
    single_item_test = [{
        'part_number': 'LS2000-115VAC-S-12',
        'quantity': 2,
        'type': 'main',
        'data': {
            'total_price': 1250.00,
            'voltage': '115VAC',
            'probe_length': 12,
            'supply_voltage': '115VAC',
            'unit_price': '1,250.00',
            'pc_size': '¾"',
            'pc_type': 'NPT',
            'pc_matt': 'SS',
            'max_pressure': '300 PSI',
            'ins_material': 'UHMWPE',
            'ins_length': '4"',
            'ins_long': 'Long',
            'ins_temp': '450',
            'probe_size': '½',
            'probe_material': '316SS'
        }
    }]
    
    success = generate_unified_quote(
        quote_items=single_item_test,
        customer_name="ACME Industrial Solutions",
        attention_name="John Smith, P.E.",
        quote_number="Q-2024-UNIFIED-001",
        output_path="test_unified_single.docx"
    )
    
    if success:
        print("✓ Single item unified quote generated successfully!")
    else:
        print("❌ Single item unified quote generation failed")
        return False
    
    # Test with multi-item
    multi_item_test = [
        {
            'part_number': 'LS2000-115VAC-S-12',
            'quantity': 1,
            'type': 'main',
            'data': {
                'total_price': 1250.00,
                'voltage': '115VAC',
                'probe_length': 12,
                'unit_price': '1,250.00'
            }
        },
        {
            'part_number': 'LS7000-230VAC-S-18',
            'quantity': 2,
            'type': 'main',
            'data': {
                'total_price': 1850.00,
                'voltage': '230VAC',
                'probe_length': 18,
                'unit_price': '1,850.00'
            }
        }
    ]
    
    success = generate_unified_quote(
        quote_items=multi_item_test,
        customer_name="ACME Industrial Solutions",
        attention_name="John Smith, P.E.",
        quote_number="Q-2024-UNIFIED-002",
        output_path="test_unified_multi.docx"
    )
    
    if success:
        print("✓ Multi-item unified quote generated successfully!")
        return True
    else:
        print("❌ Multi-item unified quote generation failed")
        return False


if __name__ == "__main__":
    test_unified_processing()