"""
CORRECTED Unified Template Processor - Uses Master Template File

This processor loads the existing master_template.docx and replaces variables,
instead of creating documents from scratch.
"""

import os
import json
import re
from pathlib import Path
from typing import Dict, Any, List, Optional
from datetime import datetime
import logging

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)

class UnifiedTemplateProcessor:
    """
    CORRECTED processor that loads master_template.docx and replaces variables.
    """
    
    def __init__(self, templates_dir: Optional[str] = None):
        """Initialize with path to master template."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required. Install with: pip install python-docx")
            
        if templates_dir is None:
            templates_dir = os.path.dirname(__file__)
        
        self.templates_dir = Path(templates_dir)
        self.configs_dir = self.templates_dir / 'configs'
        self.master_template_path = self.templates_dir / 'master_template.docx'
        
        if not self.master_template_path.exists():
            raise FileNotFoundError(f"Master template not found: {self.master_template_path}")
    
    def _load_model_config(self, model: str) -> Dict[str, Any]:
        """Load configuration for a specific model."""
        config_path = self.configs_dir / f'{model}_config.json'
        try:
            with open(config_path, 'r') as f:
                return json.load(f)
        except Exception as e:
            logger.warning(f"Could not load config for model {model}: {e}")
            return self._get_default_config(model)
    
    def _get_default_config(self, model: str) -> Dict[str, Any]:
        """Get default configuration when model config is not found."""
        return {
            "model": model,
            "description": f"{model} Level Switch Quote",
            "technical_specifications": [
                {"label": "Supply Voltage", "value": "{{supply_voltage}}", "static": False},
                {"label": "Output", "value": "10 Amp SPDT Relay", "static": True},
                {"label": "Process Connection", "value": "{{pc_size}} {{pc_type}}", "static": False},
                {"label": "Probe", "value": "{{probe_size}}\" Diameter {{probe_material}} x {{probe_length}}\"", "static": False}
            ],
            "optional_sections": {"notes": {"enabled": False, "content": []}},
            "default_values": {}
        }
    
    def _extract_model_from_part_number(self, part_number: str) -> str:
        """Extract model from part number."""
        if '-' in part_number:
            return part_number.split('-')[0]
        elif len(part_number) >= 6:
            return part_number[:6]
        else:
            return part_number
    
    def _build_items_section(self, quote_items: List[Dict[str, Any]], base_variables: Dict[str, str]) -> str:
        """Build items section content matching original template format."""
        if not quote_items:
            return ""
        
        if len(quote_items) == 1:
            return self._build_single_item_section(quote_items[0], base_variables)
        else:
            return self._build_multi_item_section(quote_items, base_variables)
    
    def _build_single_item_section(self, item: Dict[str, Any], base_variables: Dict[str, str]) -> str:
        """Build single item section exactly like original LS2000 template."""
        part_number = item.get('part_number', '')
        quantity = item.get('quantity', 1)
        
        # Get model and config
        model = self._extract_model_from_part_number(part_number)
        config = self._load_model_config(model)
        
        # Merge item data with base variables
        item_variables = base_variables.copy()
        if 'data' in item:
            item_variables.update(item['data'])
        item_variables.update({
            'part_number': part_number,
            'quantity': str(quantity),
            'model': model
        })
        
        # Format unit price
        if item.get('type') == 'main':
            unit_price = item.get('data', {}).get('total_price', 0.0)
        else:
            unit_price = item.get('data', {}).get('pricing', {}).get('total_price', 0.0)
        
        price_str = item_variables.get('unit_price', f"{unit_price:.2f}")
        if not price_str.startswith('$') and not price_str.startswith('Price'):
            price_str = f"${price_str}"
        
        # Build the section exactly like original template
        section = f"{quantity} QTY {part_number}             {price_str}    EACH"
        
        # Add technical specifications as bullet points
        specs = config.get('technical_specifications', [])
        for spec in specs:
            label = spec.get('label', '')
            value_template = spec.get('value', '')
            
            # Replace variables in the value
            value = value_template
            for var_name, var_value in item_variables.items():
                placeholder = f"{{{{{var_name}}}}}"
                if var_value:
                    value = value.replace(placeholder, str(var_value))
            
            # Clean up any remaining placeholders
            value = re.sub(r'\{\{[^}]+\}\}', '', value).strip()
            
            if value:
                section += f"\n• {label}: {value}"
        
        return section
    
    def _build_multi_item_section(self, quote_items: List[Dict[str, Any]], base_variables: Dict[str, str]) -> str:
        """Build multi-item section with clean formatting."""
        sections = []
        
        for i, item in enumerate(quote_items, 1):
            part_number = item.get('part_number', '')
            quantity = item.get('quantity', 1)
            
            # Get model and config
            model = self._extract_model_from_part_number(part_number)
            config = self._load_model_config(model)
            
            # Merge item data
            item_variables = base_variables.copy()
            if 'data' in item:
                item_variables.update(item['data'])
            item_variables.update({
                'part_number': part_number,
                'quantity': str(quantity),
                'model': model
            })
            
            # Get price
            if item.get('type') == 'main':
                unit_price = item.get('data', {}).get('total_price', 0.0)
            else:
                unit_price = item.get('data', {}).get('pricing', {}).get('total_price', 0.0)
            
            # Build item section
            item_section = f"Item {i}: {quantity} QTY {part_number}             ${unit_price:.2f} EACH"
            
            # Add specs for this item
            specs = config.get('technical_specifications', [])
            for spec in specs:
                label = spec.get('label', '')
                value_template = spec.get('value', '')
                
                value = value_template
                for var_name, var_value in item_variables.items():
                    placeholder = f"{{{{{var_name}}}}}"
                    if var_value:
                        value = value.replace(placeholder, str(var_value))
                
                value = re.sub(r'\{\{[^}]+\}\}', '', value).strip()
                
                if value:
                    item_section += f"\n• {label}: {value}"
            
            sections.append(item_section)
        
        return "\n\n".join(sections)
    
    def _build_quote_summary_table(self, quote_items: List[Dict[str, Any]]) -> str:
        """Build quote summary for multi-item quotes."""
        if len(quote_items) <= 1:
            return ""
        
        total = 0.0
        for item in quote_items:
            quantity = item.get('quantity', 1)
            if item.get('type') == 'main':
                unit_price = item.get('data', {}).get('total_price', 0.0)
            else:
                unit_price = item.get('data', {}).get('pricing', {}).get('total_price', 0.0)
            total += unit_price * quantity
        
        return f"Quote Total: ${total:.2f}"
    
    def _build_optional_notes_section(self, quote_items: List[Dict[str, Any]]) -> str:
        """Build optional notes section."""
        notes = []
        processed_models = set()
        
        for item in quote_items:
            part_number = item.get('part_number', '')
            model = self._extract_model_from_part_number(part_number)
            
            if model in processed_models:
                continue
            processed_models.add(model)
            
            config = self._load_model_config(model)
            notes_config = config.get('optional_sections', {}).get('notes', {})
            
            if notes_config.get('enabled', False):
                content = notes_config.get('content', [])
                for note in content:
                    notes.append(f"• {note}")
        
        return "\n".join(notes) if notes else ""
    
    def process_unified_template(self, quote_items: List[Dict[str, Any]], variables: Dict[str, Any]) -> Optional[Any]:
        """
        CORRECTED: Load master template and replace variables.
        
        Args:
            quote_items: List of quote items
            variables: Base template variables
            
        Returns:
            Processed Document object or None if failed
        """
        try:
            logger.info(f"Loading master template: {self.master_template_path}")
            
            # CRITICAL: Load the existing master template instead of creating new document
            doc = Document(str(self.master_template_path))
            
            # Prepare all variable replacements
            str_variables = {k: str(v) if v is not None else "" for k, v in variables.items()}
            
            # Determine quote type
            is_multi_item = len(quote_items) > 1
            str_variables['has_multiple_items'] = str(is_multi_item).lower()
            str_variables['item_count'] = str(len(quote_items))
            
            # Build dynamic content sections
            if is_multi_item:
                str_variables['model_description'] = "Multi-Item Quote"
            else:
                model = self._extract_model_from_part_number(quote_items[0].get('part_number', ''))
                config = self._load_model_config(model)
                str_variables['model_description'] = config.get('description', f"{model} Level Switch Quote")
            
            # Build the main content sections
            str_variables['items_section'] = self._build_items_section(quote_items, str_variables)
            str_variables['quote_summary_table'] = self._build_quote_summary_table(quote_items)
            str_variables['optional_notes_section'] = self._build_optional_notes_section(quote_items)
            
            # Process conditional content first
            self._process_conditional_content(doc, str_variables)
            
            # Replace all variables in paragraphs
            for paragraph in doc.paragraphs:
                self._replace_variables_in_paragraph(paragraph, str_variables)
            
            # Replace variables in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._replace_variables_in_paragraph(paragraph, str_variables)
            
            # Process headers and footers
            for section in doc.sections:
                for paragraph in section.header.paragraphs:
                    self._replace_variables_in_paragraph(paragraph, str_variables)
                for paragraph in section.footer.paragraphs:
                    self._replace_variables_in_paragraph(paragraph, str_variables)
            
            logger.info("Template processing completed successfully")
            return doc
            
        except Exception as e:
            logger.error(f"Error processing master template: {e}")
            import traceback
            logger.error(f"Traceback: {traceback.format_exc()}")
            return None
    
    def _replace_variables_in_paragraph(self, paragraph, variables: Dict[str, str]) -> None:
        """Replace variables in a paragraph, handling multi-line content."""
        full_text = paragraph.text
        
        # Find variables that need replacement
        variable_pattern = r'\{\{([^}]+)\}\}'
        matches = re.findall(variable_pattern, full_text)
        
        if not matches:
            return
        
        new_text = full_text
        for var_name in matches:
            placeholder = f"{{{{{var_name}}}}}"
            value = variables.get(var_name, f"{{{{MISSING: {var_name}}}}}")
            new_text = new_text.replace(placeholder, str(value))
        
        # Replace paragraph text
        paragraph.clear()
        
        # Handle multi-line content
        lines = new_text.split('\n')
        for i, line in enumerate(lines):
            if i > 0:
                paragraph.add_run('\n')
            paragraph.add_run(line)
    
    def _process_conditional_content(self, doc, variables: Dict[str, str]) -> None:
        """Process conditional content markers."""
        is_multi_item = variables.get('has_multiple_items', 'false').lower() == 'true'
        
        for paragraph in doc.paragraphs:
            self._process_conditional_paragraph(paragraph, is_multi_item)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._process_conditional_paragraph(paragraph, is_multi_item)
    
    def _process_conditional_paragraph(self, paragraph, is_multi_item: bool) -> None:
        """Process conditional content in a paragraph."""
        full_text = paragraph.text
        
        # Process single item conditionals
        single_pattern = r'\{\{if_single_item:([^}]+)\}\}'
        multi_pattern = r'\{\{if_multiple_items:([^}]+)\}\}'
        
        # Handle single item conditionals
        if not is_multi_item:
            full_text = re.sub(single_pattern, r'\1', full_text)
            full_text = re.sub(multi_pattern, '', full_text)
        else:
            full_text = re.sub(single_pattern, '', full_text)
            full_text = re.sub(multi_pattern, r'\1', full_text)
        
        if full_text != paragraph.text:
            paragraph.clear()
            paragraph.add_run(full_text)
    
    def save_document(self, doc: Any, output_path: str) -> bool:
        """Save the processed document."""
        try:
            doc.save(output_path)
            logger.info(f"Document saved successfully to: {output_path}")
            return True
        except Exception as e:
            logger.error(f"Error saving document to {output_path}: {e}")
            return False


# Convenience function matching the original API
def generate_unified_quote(
    quote_items: List[Dict[str, Any]],
    customer_name: str,
    attention_name: str,
    quote_number: str,
    output_path: str,
    employee_info: Optional[Dict[str, str]] = None,
    **kwargs
) -> bool:
    """Generate quote using the CORRECTED unified template system."""
    if not DOCX_AVAILABLE:
        logger.error("python-docx not available")
        return False
    
    try:
        processor = UnifiedTemplateProcessor()
        
        # Get employee info
        employee_name = kwargs.get('employee_name', 'John Nicholosi')
        employee_phone = kwargs.get('employee_phone', '(713) 467-4438')
        employee_email = kwargs.get('employee_email', 'John@babbitt.us')
        
        if employee_info:
            employee_name = employee_info.get('name', employee_name)
            employee_phone = employee_info.get('phone', employee_phone)
            employee_email = employee_info.get('email', employee_email)
        
        # Prepare variables
        variables = {
            'date': datetime.now().strftime("%B %d, %Y"),
            'customer_name': customer_name,
            'attention_name': attention_name,
            'quote_number': quote_number,
            'employee_name': employee_name,
            'employee_phone': employee_phone,
            'employee_email': employee_email,
            'lead_time': kwargs.get('lead_time', 'In Stock')
        }
        
        variables.update(kwargs)
        
        # Process the template
        doc = processor.process_unified_template(quote_items, variables)
        if not doc:
            logger.error("Failed to process unified template")
            return False
        
        # Save the document
        return processor.save_document(doc, output_path)
        
    except Exception as e:
        logger.error(f"Error generating unified quote: {e}")
        return False