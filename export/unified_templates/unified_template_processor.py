"""
CORRECTED Unified Template Processor - Uses Master Template File

This processor loads the existing master_template.docx and replaces variables,
instead of creating documents from scratch.
"""

import os
import json
import re
from pathlib import Path
from typing import Dict, Any, List, Optional
from datetime import datetime
import logging

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)

class UnifiedTemplateProcessor:
    """
    CORRECTED processor that loads master_template.docx and replaces variables.
    """
    
    def __init__(self, templates_dir: Optional[str] = None):
        """Initialize with path to master template."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required. Install with: pip install python-docx")
            
        if templates_dir is None:
            templates_dir = os.path.dirname(__file__)
        
        self.templates_dir = Path(templates_dir)
        self.configs_dir = self.templates_dir / 'configs'
        self.master_template_path = self.templates_dir / 'master_template.docx'
        
        # Check for model-specific templates
        self.model_templates_dir = self.templates_dir.parent / 'templates'
        
        if not self.master_template_path.exists():
            raise FileNotFoundError(f"Master template not found: {self.master_template_path}")
    
    def _load_model_config(self, model: str) -> Dict[str, Any]:
        """Load configuration for a specific model."""
        config_path = self.configs_dir / f'{model}_config.json'
        try:
            with open(config_path, 'r') as f:
                return json.load(f)
        except Exception as e:
            logger.warning(f"Could not load config for model {model}: {e}")
            return self._get_default_config(model)
    
    def _get_default_config(self, model: str) -> Dict[str, Any]:
        """Get default configuration when model config is not found."""
        return {
            "model": model,
            "description": f"{model} Level Switch Quote",
            "technical_specifications": [
                {"label": "Supply Voltage", "value": "{{supply_voltage}}", "static": False},
                {"label": "Output", "value": "10 Amp SPDT Relay", "static": True},
                {"label": "Process Connection", "value": "{{pc_size}} {{pc_type}}", "static": False},
                {"label": "Probe", "value": "{{probe_size}}\" Diameter {{probe_material}} x {{probe_length}}\"", "static": False}
            ],
            "optional_sections": {"notes": {"enabled": False, "content": []}},
            "default_values": {}
        }
    
    def _extract_model_from_part_number(self, part_number: str) -> str:
        """Extract model from part number."""
        if '-' in part_number:
            return part_number.split('-')[0]
        elif len(part_number) >= 6:
            return part_number[:6]
        else:
            return part_number
    
    def _get_template_path(self, model: str) -> Optional[Path]:
        """Get the appropriate template path for a model."""
        # First check for model-specific template
        model_template_path = self.model_templates_dir / f'{model}_template.docx'
        if model_template_path.exists():
            return model_template_path
        
        # Fallback to master template
        if self.master_template_path.exists():
            return self.master_template_path
        
        return None
    
    def _build_items_section(self, quote_items: List[Dict[str, Any]], base_variables: Dict[str, str]) -> str:
        """Build items section content matching original template format."""
        if not quote_items:
            return ""
        
        if len(quote_items) == 1:
            return self._build_single_item_section(quote_items[0], base_variables)
        else:
            return self._build_multi_item_section(quote_items, base_variables)
    
    def _build_single_item_section(self, item: Dict[str, Any], base_variables: Dict[str, str]) -> str:
        """Build single item section exactly like original LS2000 template."""
        part_number = item.get('part_number', '')
        quantity = item.get('quantity', 1)
        
        # Get model and config
        model = self._extract_model_from_part_number(part_number)
        config = self._load_model_config(model)
        
        # Merge item data with base variables
        item_variables = base_variables.copy()
        if 'data' in item:
            # Map quote data to template variables
            mapped_vars = self._map_quote_data_to_template_variables(item['data'])
            item_variables.update(mapped_vars)
            # Also add original data for backward compatibility
            item_variables.update(item['data'])
        item_variables.update({
            'part_number': part_number,
            'quantity': str(quantity),
            'model': model
        })
        
        # Format unit price
        if item.get('type') == 'main':
            unit_price = item.get('data', {}).get('total_price', 0.0)
        else:
            unit_price = item.get('data', {}).get('pricing', {}).get('total_price', 0.0)
        
        price_str = item_variables.get('unit_price', f"{unit_price:.2f}")
        if not price_str.startswith('$') and not price_str.startswith('Price'):
            price_str = f"${price_str}"
        
        # Build the section exactly like original template
        section = f"{quantity} QTY {part_number}             {price_str}    EACH"
        
        # Add technical specifications as bullet points
        specs = config.get('technical_specifications', [])
        for spec in specs:
            label = spec.get('label', '')
            value_template = spec.get('value', '')
            
            # Replace variables in the value
            value = value_template
            for var_name, var_value in item_variables.items():
                placeholder = f"{{{{{var_name}}}}}"
                if var_value:
                    value = value.replace(placeholder, str(var_value))
            
            # Clean up any remaining placeholders
            value = re.sub(r'\{\{[^}]+\}\}', '', value).strip()
            
            if value:
                section += f"\n• {label}: {value}"
        
        return section
    
    def _build_multi_item_section(self, quote_items: List[Dict[str, Any]], base_variables: Dict[str, str]) -> str:
        """Build multi-item section with bullet points extracted from model-specific templates."""
        sections = []
        
        for i, item in enumerate(quote_items, 1):
            part_number = item.get('part_number', '')
            quantity = item.get('quantity', 1)
            
            # Get model
            model = self._extract_model_from_part_number(part_number)
            
            # Merge item data
            item_variables = base_variables.copy()
            if 'data' in item:
                # Map quote data to template variables
                mapped_vars = self._map_quote_data_to_template_variables(item['data'])
                item_variables.update(mapped_vars)
                # Also add original data for backward compatibility
                item_variables.update(item['data'])
            item_variables.update({
                'part_number': part_number,
                'quantity': str(quantity),
                'model': model
            })
            
            # Ensure all required variables are available for template processing
            if 'attention_name' not in item_variables:
                item_variables['attention_name'] = base_variables.get('attention_name', '')
            if 'customer_name' not in item_variables:
                item_variables['customer_name'] = base_variables.get('customer_name', '')
            if 'date' not in item_variables:
                item_variables['date'] = base_variables.get('date', '')
            if 'quote_number' not in item_variables:
                item_variables['quote_number'] = base_variables.get('quote_number', '')
            if 'employee_name' not in item_variables:
                item_variables['employee_name'] = base_variables.get('employee_name', '')
            if 'employee_phone' not in item_variables:
                item_variables['employee_phone'] = base_variables.get('employee_phone', '')
            if 'employee_email' not in item_variables:
                item_variables['employee_email'] = base_variables.get('employee_email', '')
            if 'lead_time' not in item_variables:
                item_variables['lead_time'] = base_variables.get('lead_time', 'In Stock')
            
            # Ensure probe_size is available
            if 'probe_diameter' in item['data'] and 'probe_size' not in item_variables:
                probe_size = str(item['data']['probe_diameter']).replace('"', '').replace("'", '')
                item_variables['probe_size'] = probe_size
            
            # Get price
            if item.get('type') == 'main':
                unit_price = item.get('data', {}).get('total_price', 0.0)
            else:
                unit_price = item.get('data', {}).get('pricing', {}).get('total_price', 0.0)
            
            # Build item section header
            item_section = f"Item {i}: {quantity} QTY {part_number}             ${unit_price:.2f} EACH"
            
            # Extract bullet points from the model-specific template
            bullet_points = self._extract_bullet_points_from_template(model, item_variables)
            
            # Add bullet points to the item section
            for bullet_point in bullet_points:
                item_section += f"\n• {bullet_point}"
            
            sections.append(item_section)
        
        return "\n\n".join(sections)
    
    def _build_quote_summary_table(self, quote_items: List[Dict[str, Any]]) -> str:
        """Build quote summary for multi-item quotes."""
        if len(quote_items) <= 1:
            return ""
        
        total = 0.0
        for item in quote_items:
            quantity = item.get('quantity', 1)
            if item.get('type') == 'main':
                unit_price = item.get('data', {}).get('total_price', 0.0)
            else:
                unit_price = item.get('data', {}).get('pricing', {}).get('total_price', 0.0)
            total += unit_price * quantity
        
        return f"Quote Total: ${total:.2f}"
    
    def _build_optional_notes_section(self, quote_items: List[Dict[str, Any]]) -> str:
        """Build optional notes section."""
        notes = []
        processed_models = set()
        
        for item in quote_items:
            part_number = item.get('part_number', '')
            model = self._extract_model_from_part_number(part_number)
            
            if model in processed_models:
                continue
            processed_models.add(model)
            
            config = self._load_model_config(model)
            notes_config = config.get('optional_sections', {}).get('notes', {})
            
            if notes_config.get('enabled', False):
                content = notes_config.get('content', [])
                for note in content:
                    notes.append(f"• {note}")
        
        return "\n".join(notes) if notes else ""
    
    def process_unified_template(self, quote_items: List[Dict[str, Any]], variables: Dict[str, Any]) -> Optional[Any]:
        """
        CORRECTED: Load appropriate template and replace variables.
        
        For single items: Use model-specific template (e.g., LS2000_template.docx)
        For multi-items: Use master_template.docx
        
        Args:
            quote_items: List of quote items
            variables: Base template variables
            
        Returns:
            Processed Document object or None if failed
        """
        try:
            # Determine which template to use
            if len(quote_items) == 1:
                # Single item: Use model-specific template
                model = self._extract_model_from_part_number(quote_items[0].get('part_number', ''))
                template_path = self._get_template_path(model)
                if template_path and template_path.exists():
                    logger.info(f"Loading model-specific template: {template_path}")
                    doc = Document(str(template_path))
                else:
                    logger.warning(f"Model-specific template not found for {model}, using master template")
                    doc = Document(str(self.master_template_path))
            else:
                # Multi-item: Use master template
                logger.info(f"Loading master template: {self.master_template_path}")
                doc = Document(str(self.master_template_path))
            
            # Prepare all variable replacements
            str_variables = {k: str(v) if v is not None else "" for k, v in variables.items()}
            
            # For single items, also include the original quote data directly
            if len(quote_items) == 1:
                item_data = quote_items[0].get('data', {})
                # Add original data variables directly
                for key, value in item_data.items():
                    if value is not None:
                        str_variables[key] = str(value)
                
                # Also add mapped variables for backward compatibility
                mapped_vars = self._map_quote_data_to_template_variables(item_data)
                for key, value in mapped_vars.items():
                    str_variables[key] = value
                
                # Ensure unit_price is available
                if 'total_price' in item_data and 'unit_price' not in str_variables:
                    str_variables['unit_price'] = f"${item_data['total_price']:.2f}"
                
                # Ensure pc_rate is available (even if None)
                if 'pc_rate' in item_data:
                    if item_data['pc_rate'] is None:
                        str_variables['pc_rate'] = ""
                    else:
                        str_variables['pc_rate'] = str(item_data['pc_rate'])
                else:
                    str_variables['pc_rate'] = ""
                
                # Ensure quantity is available
                if 'quantity' not in str_variables:
                    str_variables['quantity'] = str(quote_items[0].get('quantity', 1))
                
                # Ensure all base variables are available
                if 'attention_name' not in str_variables:
                    str_variables['attention_name'] = variables.get('attention_name', '')
                if 'customer_name' not in str_variables:
                    str_variables['customer_name'] = variables.get('customer_name', '')
                if 'date' not in str_variables:
                    str_variables['date'] = variables.get('date', '')
                if 'quote_number' not in str_variables:
                    str_variables['quote_number'] = variables.get('quote_number', '')
                if 'employee_name' not in str_variables:
                    str_variables['employee_name'] = variables.get('employee_name', '')
                if 'employee_phone' not in str_variables:
                    str_variables['employee_phone'] = variables.get('employee_phone', '')
                if 'employee_email' not in str_variables:
                    str_variables['employee_email'] = variables.get('employee_email', '')
                if 'lead_time' not in str_variables:
                    str_variables['lead_time'] = variables.get('lead_time', 'In Stock')
            
            # Determine quote type
            is_multi_item = len(quote_items) > 1
            str_variables['has_multiple_items'] = str(is_multi_item).lower()
            str_variables['item_count'] = str(len(quote_items))
            
            # Build dynamic content sections
            if is_multi_item:
                str_variables['model_description'] = "Multi-Item Quote"
            else:
                model = self._extract_model_from_part_number(quote_items[0].get('part_number', ''))
                config = self._load_model_config(model)
                str_variables['model_description'] = config.get('description', f"{model} Level Switch Quote")
            
            # Build the main content sections
            str_variables['items_section'] = self._build_items_section(quote_items, str_variables)
            str_variables['quote_summary_table'] = self._build_quote_summary_table(quote_items)
            str_variables['optional_notes_section'] = self._build_optional_notes_section(quote_items)
            
            # Process conditional content first
            self._process_conditional_content(doc, str_variables)
            
            # Replace all variables in paragraphs
            for paragraph in doc.paragraphs:
                self._replace_variables_in_paragraph(paragraph, str_variables)
            
            # Replace variables in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._replace_variables_in_paragraph(paragraph, str_variables)
            
            # Process headers and footers
            for section in doc.sections:
                for paragraph in section.header.paragraphs:
                    self._replace_variables_in_paragraph(paragraph, str_variables)
                for paragraph in section.footer.paragraphs:
                    self._replace_variables_in_paragraph(paragraph, str_variables)
            
            logger.info("Template processing completed successfully")
            return doc
            
        except Exception as e:
            logger.error(f"Error processing template: {e}")
            import traceback
            logger.error(f"Traceback: {traceback.format_exc()}")
            return None
    
    def _replace_variables_in_paragraph(self, paragraph, variables: Dict[str, str]) -> None:
        """Replace variables in a paragraph, handling multi-line content."""
        full_text = paragraph.text
        
        # Find variables that need replacement
        variable_pattern = r'\{\{([^}]+)\}\}'
        matches = re.findall(variable_pattern, full_text)
        
        if not matches:
            return
        
        new_text = full_text
        for var_name in matches:
            placeholder = f"{{{{{var_name}}}}}"
            value = variables.get(var_name, f"{{{{MISSING: {var_name}}}}}")
            new_text = new_text.replace(placeholder, str(value))
        
        # Replace paragraph text
        paragraph.clear()
        
        # Handle multi-line content
        lines = new_text.split('\n')
        for i, line in enumerate(lines):
            if i > 0:
                paragraph.add_run('\n')
            paragraph.add_run(line)
    
    def _process_conditional_content(self, doc, variables: Dict[str, str]) -> None:
        """Process conditional content markers."""
        is_multi_item = variables.get('has_multiple_items', 'false').lower() == 'true'
        
        for paragraph in doc.paragraphs:
            self._process_conditional_paragraph(paragraph, is_multi_item)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._process_conditional_paragraph(paragraph, is_multi_item)
    
    def _process_conditional_paragraph(self, paragraph, is_multi_item: bool) -> None:
        """Process conditional content in a paragraph."""
        full_text = paragraph.text
        
        # Process single item conditionals
        single_pattern = r'\{\{if_single_item:([^}]+)\}\}'
        multi_pattern = r'\{\{if_multiple_items:([^}]+)\}\}'
        
        # Handle single item conditionals
        if not is_multi_item:
            full_text = re.sub(single_pattern, r'\1', full_text)
            full_text = re.sub(multi_pattern, '', full_text)
        else:
            full_text = re.sub(single_pattern, '', full_text)
            full_text = re.sub(multi_pattern, r'\1', full_text)
        
        if full_text != paragraph.text:
            paragraph.clear()
            paragraph.add_run(full_text)
    
    def save_document(self, doc: Any, output_path: str) -> bool:
        """Save the processed document."""
        try:
            doc.save(output_path)
            logger.info(f"Document saved successfully to: {output_path}")
            return True
        except Exception as e:
            logger.error(f"Error saving document to {output_path}: {e}")
            return False

    def _map_quote_data_to_template_variables(self, item_data: Dict[str, Any]) -> Dict[str, str]:
        """Map quote data fields to template variable names."""
        mapped_vars = {}
        
        # Map voltage fields
        if 'voltage' in item_data:
            mapped_vars['supply_voltage'] = str(item_data['voltage'])
        
        # Map insulator fields
        if 'insulator' in item_data:
            insulator_str = str(item_data['insulator'])
            # Parse insulator string like "2.0" UHMWPE" or "4.0" Teflon"
            if '"' in insulator_str:
                parts = insulator_str.split('"')
                if len(parts) >= 2:
                    # Extract length and remove decimal places
                    length_raw = parts[0].strip()
                    try:
                        # Convert to float and back to remove .0
                        length_float = float(length_raw)
                        length = str(int(length_float)) if length_float.is_integer() else str(length_float)
                    except ValueError:
                        length = length_raw
                    
                    material = parts[1].strip()
                    mapped_vars['ins_material'] = material
                    mapped_vars['ins_length'] = length
                    # Create the formatted long description without duplication
                    mapped_vars['ins_long'] = f"{length}\" {material}"
                    # Set temperature rating based on material
                    if 'Teflon' in material:
                        mapped_vars['ins_temp'] = '450'
                    else:
                        mapped_vars['ins_temp'] = '180'
        
        # Map probe fields
        if 'probe_material_name' in item_data:
            mapped_vars['probe_material'] = str(item_data['probe_material_name'])
        elif 'probe_material' in item_data:
            mapped_vars['probe_material'] = str(item_data['probe_material'])
        
        if 'probe_diameter' in item_data:
            probe_size = str(item_data['probe_diameter']).replace('"', '').replace("'", '')
            mapped_vars['probe_size'] = probe_size
        elif 'probe_size' in item_data:
            mapped_vars['probe_size'] = str(item_data['probe_size'])
        
        if 'probe_length' in item_data:
            probe_length_value = item_data['probe_length']
            if isinstance(probe_length_value, (int, float)):
                # Remove .0 if it's a whole number, otherwise keep the decimal
                if probe_length_value.is_integer() if hasattr(probe_length_value, 'is_integer') else probe_length_value == int(probe_length_value):
                    mapped_vars['probe_length'] = str(int(probe_length_value))
                else:
                    mapped_vars['probe_length'] = str(probe_length_value)
            else:
                mapped_vars['probe_length'] = str(probe_length_value)
        
        # Map process connection fields
        if 'pc_size' in item_data:
            mapped_vars['pc_size'] = str(item_data['pc_size'])
        
        if 'pc_type' in item_data:
            mapped_vars['pc_type'] = str(item_data['pc_type'])
        
        if 'pc_matt' in item_data:
            mapped_vars['pc_matt'] = str(item_data['pc_matt'])
        
        # Handle pc_rate - set to empty string if None, otherwise use the value
        if 'pc_rate' in item_data:
            if item_data['pc_rate'] is None:
                mapped_vars['pc_rate'] = ""
            else:
                mapped_vars['pc_rate'] = str(item_data['pc_rate'])
        
        # Map pressure and temperature
        if 'max_pressure' in item_data and item_data['max_pressure']:
            if isinstance(item_data['max_pressure'], (int, float)):
                mapped_vars['max_pressure'] = f"{item_data['max_pressure']} PSI"
            else:
                mapped_vars['max_pressure'] = str(item_data['max_pressure'])
        
        if 'max_temperature' in item_data and item_data['max_temperature']:
            if isinstance(item_data['max_temperature'], (int, float)):
                mapped_vars['max_temperature'] = f"{item_data['max_temperature']}°F"
            else:
                mapped_vars['max_temperature'] = str(item_data['max_temperature'])
        
        # Map unit_price from total_price
        if 'total_price' in item_data:
            mapped_vars['unit_price'] = f"${item_data['total_price']:.2f}"
        
        return mapped_vars

    def _extract_bullet_points_from_template(self, model: str, item_variables: Dict[str, str]) -> List[str]:
        """
        Extract bullet points from a model-specific template by processing it with the item's variables.
        This simulates what the single-item template would show for this specific item.
        """
        try:
            # Get the model-specific template path
            template_path = self._get_template_path(model)
            if not template_path or not template_path.exists():
                logger.warning(f"Model-specific template not found for {model}, using config fallback")
                return self._get_bullet_points_from_config(model, item_variables)
            
            # Ensure all required variables are available
            processed_variables = item_variables.copy()
            
            # Ensure probe_size is available if probe_diameter exists
            if 'probe_diameter' in processed_variables and 'probe_size' not in processed_variables:
                probe_size = str(processed_variables['probe_diameter']).replace('"', '').replace("'", '')
                processed_variables['probe_size'] = probe_size
            
            # Ensure unit_price is available if total_price exists
            if 'total_price' in processed_variables and 'unit_price' not in processed_variables:
                processed_variables['unit_price'] = f"${float(processed_variables['total_price']):.2f}"
            
            # Ensure pc_rate is available
            if 'pc_rate' not in processed_variables:
                processed_variables['pc_rate'] = ""
            
            # Load the template
            doc = Document(str(template_path))
            
            # Process the template with the item's variables
            processed_lines = []
            
            # Debug: Print what variables we're using
            logger.debug(f"Processing template for {model} with variables: {list(processed_variables.keys())}")
            
            # Process paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    original_text = paragraph.text.strip()
                    processed_text = self._replace_variables_in_text(original_text, processed_variables)
                    if processed_text.strip():
                        processed_lines.append(processed_text.strip())
                        # Debug: Show what was processed
                        if 'probe' in original_text.lower():
                            logger.debug(f"Probe line - Original: '{original_text}' -> Processed: '{processed_text}'")
            
            # Process tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if paragraph.text.strip():
                                original_text = paragraph.text.strip()
                                processed_text = self._replace_variables_in_text(original_text, processed_variables)
                                if processed_text.strip():
                                    processed_lines.append(processed_text.strip())
                                    # Debug: Show what was processed
                                    if 'probe' in original_text.lower():
                                        logger.debug(f"Probe line (table) - Original: '{original_text}' -> Processed: '{processed_text}'")
            
            # Extract bullet points (lines starting with • or -)
            bullet_points = []
            for line in processed_lines:
                line = line.strip()
                if line.startswith('•') or line.startswith('-') or line.startswith('*'):
                    # Remove the bullet and clean up
                    bullet_text = line[1:].strip()
                    if bullet_text:
                        bullet_points.append(bullet_text)
                elif ':' in line and any(keyword in line.lower() for keyword in ['voltage', 'output', 'connection', 'insulator', 'probe', 'housing', 'warranty']):
                    # This looks like a technical specification line
                    bullet_points.append(line)
            
            return bullet_points
            
        except Exception as e:
            logger.error(f"Error extracting bullet points from template for {model}: {e}")
            return self._get_bullet_points_from_config(model, item_variables)
    
    def _get_bullet_points_from_config(self, model: str, item_variables: Dict[str, str]) -> List[str]:
        """Fallback method to get bullet points from config when template extraction fails."""
        config = self._load_model_config(model)
        specs = config.get('technical_specifications', [])
        
        bullet_points = []
        for spec in specs:
            label = spec.get('label', '')
            value_template = spec.get('value', '')
            
            value = value_template
            for var_name, var_value in item_variables.items():
                placeholder = f"{{{{{var_name}}}}}"
                if var_value:
                    value = value.replace(placeholder, str(var_value))
            
            value = re.sub(r'\{\{[^}]+\}\}', '', value).strip()
            
            if value:
                bullet_points.append(f"{label}: {value}")
        
        return bullet_points
    
    def _replace_variables_in_text(self, text: str, variables: Dict[str, str]) -> str:
        """Replace variables in text without modifying paragraph objects."""
        new_text = text
        
        # Find variables that need replacement
        variable_pattern = r'\{\{([^}]+)\}\}'
        matches = re.findall(variable_pattern, text)
        
        for var_name in matches:
            placeholder = f"{{{{{var_name}}}}}"
            value = variables.get(var_name, f"{{{{MISSING: {var_name}}}}}")
            new_text = new_text.replace(placeholder, str(value))
        
        return new_text


# Convenience function matching the original API
def generate_unified_quote(
    quote_items: List[Dict[str, Any]],
    customer_name: str,
    attention_name: str,
    quote_number: str,
    output_path: str,
    employee_info: Optional[Dict[str, str]] = None,
    **kwargs
) -> bool:
    """Generate quote using the CORRECTED unified template system."""
    if not DOCX_AVAILABLE:
        logger.error("python-docx not available")
        return False
    
    try:
        processor = UnifiedTemplateProcessor()
        
        # Get employee info
        employee_name = kwargs.get('employee_name', 'John Nicholosi')
        employee_phone = kwargs.get('employee_phone', '(713) 467-4438')
        employee_email = kwargs.get('employee_email', 'John@babbitt.us')
        
        if employee_info:
            employee_name = employee_info.get('name', employee_name)
            employee_phone = employee_info.get('phone', employee_phone)
            employee_email = employee_info.get('email', employee_email)
        
        # Prepare variables
        variables = {
            'date': datetime.now().strftime("%B %d, %Y"),
            'customer_name': customer_name,
            'attention_name': attention_name,
            'quote_number': quote_number,
            'employee_name': employee_name,
            'employee_phone': employee_phone,
            'employee_email': employee_email,
            'lead_time': kwargs.get('lead_time', 'In Stock')
        }
        
        variables.update(kwargs)
        
        # Process the template
        doc = processor.process_unified_template(quote_items, variables)
        if not doc:
            logger.error("Failed to process unified template")
            return False
        
        # Save the document
        return processor.save_document(doc, output_path)
        
    except Exception as e:
        logger.error(f"Error generating unified quote: {e}")
        return False