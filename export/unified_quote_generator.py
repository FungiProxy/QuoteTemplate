"""
Unified Quote Generation System

This module provides a consistent quote generation system that maintains 
uniform formatting for both single and multi-item quotes using template-based 
approach with dynamic content generation.
"""

import os
import re
from pathlib import Path
from typing import Dict, Any, Optional, List, TYPE_CHECKING
from datetime import datetime
import logging

if TYPE_CHECKING:
    from docx import Document

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)

class UnifiedQuoteGenerator:
    """
    Unified quote generator that creates consistent formatting for both 
    single and multi-item quotes using a template-based approach.
    """
    
    def __init__(self, templates_dir: Optional[str] = None):
        """
        Initialize the unified quote generator.
        
        Args:
            templates_dir: Path to the templates directory
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required. Install with: pip install python-docx")
            
        if templates_dir is None:
            templates_dir = os.path.join(os.path.dirname(__file__), 'templates')
        self.templates_dir = Path(templates_dir)
        
    def get_base_template_path(self, primary_model: str) -> Optional[Path]:
        """
        Get the base template path for the primary model.
        This will be used as the foundation for the quote structure.
        
        Args:
            primary_model: Primary model name (e.g., 'LS2000', 'FS10000')
            
        Returns:
            Path to the base template file, or None if not found
        """
        template_name = f"{primary_model}_template.docx"
        template_path = self.templates_dir / template_name
        
        if template_path.exists():
            return template_path
        
        # Try without the housing suffix (e.g., LS2000S -> LS2000)
        if len(primary_model) > 6:
            base_model = primary_model[:6]
            template_name = f"{base_model}_template.docx"
            template_path = self.templates_dir / template_name
            if template_path.exists():
                return template_path
        
        logger.warning(f"Template not found for model: {primary_model}")
        return None
    
    def extract_model_from_part_number(self, part_number: str) -> str:
        """
        Extract model from part number.
        
        Args:
            part_number: Full part number
            
        Returns:
            Model name (e.g., 'LS2000', 'FS10000')
        """
        if '-' in part_number:
            return part_number.split('-')[0]
        else:
            # Fallback: take first 6 characters
            return part_number[:6]
    
    def _replace_template_variables(self, doc: Any, variables: Dict[str, str]) -> None:
        """
        Replace template variables in the document while preserving formatting.
        
        Args:
            doc: Document object
            variables: Dictionary of variable name -> value mappings
        """
        # Process all paragraphs
        for paragraph in doc.paragraphs:
            self._replace_variables_in_paragraph(paragraph, variables)
        
        # Process all tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_variables_in_paragraph(paragraph, variables)
        
        # Process headers and footers
        for section in doc.sections:
            header = section.header
            for paragraph in header.paragraphs:
                self._replace_variables_in_paragraph(paragraph, variables)
            
            footer = section.footer
            for paragraph in footer.paragraphs:
                self._replace_variables_in_paragraph(paragraph, variables)
    
    def _replace_variables_in_paragraph(self, paragraph, variables: Dict[str, str]) -> None:
        """
        Replace template variables in a paragraph while preserving formatting.
        
        Args:
            paragraph: docx paragraph object
            variables: Dictionary of variable name -> value mappings
        """
        full_text = paragraph.text
        
        # Find all template variables in the format {{variable_name}}
        variable_pattern = r'\{\{([^}]+)\}\}'
        matches = re.findall(variable_pattern, full_text)
        
        if not matches:
            return
        
        # Replace variables in the full text
        new_text = full_text
        for var_name in matches:
            placeholder = f"{{{{{var_name}}}}}"
            value = variables.get(var_name, f"{{{{MISSING: {var_name}}}}}")
            new_text = new_text.replace(placeholder, str(value))
        
        # Only update if text changed
        if new_text != full_text:
            paragraph.clear()
            paragraph.add_run(new_text)
    
    def _add_multi_item_section(self, doc: Any, quote_items: List[Dict[str, Any]], 
                               start_index: int = 1) -> None:
        """
        Add additional items section to the document with consistent formatting.
        
        Args:
            doc: Document object
            quote_items: List of quote items
            start_index: Index to start from (1 for second item, etc.)
        """
        if len(quote_items) <= start_index:
            return  # No additional items to add
        
        # Add section header for additional items
        if start_index == 1:  # First additional item
            # Add a page break or significant spacing
            doc.add_paragraph().add_run().add_break()
            
            # Create header paragraph manually since template may not have Heading 2 style
            header = doc.add_paragraph("ADDITIONAL QUOTE ITEMS")
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Make it bold and larger
            header_run = header.runs[0]
            header_run.bold = True
            header_run.font.size = Pt(14)
        
        # Add each additional item
        for i in range(start_index, len(quote_items)):
            item = quote_items[i]
            item_number = i + 1
            
            # Add separator line
            if i > start_index:
                separator = doc.add_paragraph()
                separator.add_run("─" * 80)
                separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add item header - create manually to avoid style issues
            item_header = doc.add_paragraph(f"ITEM {item_number}: {item.get('type', 'MAIN').upper()}")
            header_run = item_header.runs[0]
            header_run.bold = True
            header_run.font.size = Pt(12)
            
            # Add basic item information
            info_para = doc.add_paragraph()
            info_para.add_run("Part Number: ").bold = True
            info_para.add_run(item.get('part_number', 'N/A'))
            
            info_para.add_run("\nQuantity: ").bold = True
            info_para.add_run(str(item.get('quantity', 1)))
            
            # Add item-specific details
            if item.get('type') == 'main':
                self._add_main_item_details(doc, item)
            else:
                self._add_spare_part_details(doc, item)
    
    def _add_main_item_details(self, doc: Any, item: Dict[str, Any]) -> None:
        """
        Add main item technical details with consistent formatting.
        
        Args:
            doc: Document object
            item: Quote item data
        """
        item_data = item.get('data', {})
        
        # Technical specifications header
        spec_header = doc.add_paragraph()
        spec_header.add_run("Technical Specifications:").bold = True
        
        # Create specifications list
        specs = [
            ("Model", item_data.get('model', 'N/A')),
            ("Supply Voltage", item_data.get('voltage', 'N/A')),
            ("Probe Length", f"{item_data.get('probe_length', 'N/A')}\"" if item_data.get('probe_length') else 'N/A'),
            ("Probe Material", item_data.get('probe_material_name', item_data.get('probe_material', 'N/A'))),
            ("Probe Diameter", item_data.get('probe_diameter', 'N/A')),
            ("Insulator Material", self._extract_insulator_material_name(item_data)),
            ("Insulator Length", f"{item_data.get('base_insulator_length', 4)}\""),
            ("Maximum Temperature", f"{item_data.get('max_temperature', 450)}°F"),
            ("Maximum Pressure", f"{item_data.get('max_pressure', 300)} PSI"),
            ("Process Connection", f"{item_data.get('pc_type', 'NPT')} {item_data.get('pc_size', '¾')}\""),
            ("Output Type", item_data.get('output_type', '10 Amp SPDT Relay'))
        ]
        
        # Add each specification
        for spec_name, spec_value in specs:
            if spec_value and spec_value != 'N/A':
                spec_para = doc.add_paragraph()
                spec_para.add_run(f"• {spec_name}: ").bold = True
                spec_para.add_run(str(spec_value))
        
        # Add options if present
        options = item_data.get('options', [])
        if options:
            options_para = doc.add_paragraph()
            options_para.add_run("• Options: ").bold = True
            options_text = ', '.join([opt.split(':')[0] if ':' in opt else opt for opt in options])
            options_para.add_run(options_text)
        
        # Add pricing
        pricing_para = doc.add_paragraph()
        pricing_para.add_run("• Unit Price: ").bold = True
        unit_price = item_data.get('total_price', 0)
        pricing_para.add_run(f"${unit_price:.2f}")
        
        total_price = unit_price * item.get('quantity', 1)
        pricing_para.add_run(" | Total: ").bold = True
        pricing_para.add_run(f"${total_price:.2f}")
    
    def _add_spare_part_details(self, doc: Any, item: Dict[str, Any]) -> None:
        """
        Add spare part details with consistent formatting.
        
        Args:
            doc: Document object
            item: Quote item data
        """
        item_data = item.get('data', {})
        
        # Description
        desc_para = doc.add_paragraph()
        desc_para.add_run("Description: ").bold = True
        desc_para.add_run(item_data.get('description', 'Spare Part'))
        
        # Category
        cat_para = doc.add_paragraph()
        cat_para.add_run("Category: ").bold = True
        cat_para.add_run(item_data.get('category', 'General'))
        
        # Pricing
        pricing_para = doc.add_paragraph()
        pricing_para.add_run("Unit Price: ").bold = True
        unit_price = item_data.get('pricing', {}).get('total_price', 0)
        pricing_para.add_run(f"${unit_price:.2f}")
        
        total_price = unit_price * item.get('quantity', 1)
        pricing_para.add_run(" | Total: ").bold = True
        pricing_para.add_run(f"${total_price:.2f}")
    
    def _add_quote_summary(self, doc: Any, quote_items: List[Dict[str, Any]]) -> None:
        """
        Add a professional quote summary table.
        
        Args:
            doc: Document object
            quote_items: List of all quote items
        """
        # Add summary section header - create manually to avoid style issues
        doc.add_paragraph().add_run().add_break()
        summary_header = doc.add_paragraph("QUOTE SUMMARY")
        summary_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = summary_header.runs[0]
        header_run.bold = True
        header_run.font.size = Pt(14)
        
        # Create summary table
        table = doc.add_table(rows=1, cols=5)
        try:
            table.style = 'Table Grid'
        except KeyError:
            # If Table Grid style doesn't exist, use default table style
            pass
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header row
        header_cells = table.rows[0].cells
        headers = ['Item', 'Part Number', 'Type', 'Qty', 'Total Price']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            # Make header bold
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Add data rows
        total_quote_value = 0.0
        for i, item in enumerate(quote_items, 1):
            row_cells = table.add_row().cells
            
            # Item number
            row_cells[0].text = str(i)
            
            # Part number
            row_cells[1].text = item.get('part_number', 'N/A')
            
            # Type
            row_cells[2].text = item.get('type', 'main').upper()
            
            # Quantity
            row_cells[3].text = str(item.get('quantity', 1))
            
            # Price calculation
            if item.get('type') == 'main':
                unit_price = item.get('data', {}).get('total_price', 0)
            else:  # spare part
                unit_price = item.get('data', {}).get('pricing', {}).get('total_price', 0)
            
            quantity = item.get('quantity', 1)
            item_total = unit_price * quantity
            total_quote_value += item_total
            
            row_cells[4].text = f"${item_total:.2f}"
        
        # Add total row
        total_row = table.add_row().cells
        total_row[0].text = ""
        total_row[1].text = ""
        total_row[2].text = ""
        total_row[3].text = "TOTAL:"
        total_row[4].text = f"${total_quote_value:.2f}"
        
        # Make total row bold
        for cell in total_row:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
    
    def _extract_insulator_material_name(self, quote_data: dict) -> str:
        """
        Extract insulator material name from quote data.
        
        Args:
            quote_data: Quote data dictionary
            
        Returns:
            Insulator material name
        """
        # Check for custom insulator first
        insulator = quote_data.get('insulator', {})
        if isinstance(insulator, dict) and insulator.get('material_name'):
            return insulator['material_name']
        
        # Check for insulator material field
        insulator_material = quote_data.get('insulator_material')
        if insulator_material:
            return insulator_material
        
        # Default fallback
        return 'UHMWPE'
    
    def _extract_display_quote_number(self, quote_number: str) -> str:
        """
        Extract display quote number without customer name prefix.
        
        Args:
            quote_number: Full quote number (e.g., "ACME ZF071925A")
            
        Returns:
            Display quote number without customer name (e.g., "ZF071925A")
        """
        # Split by space to separate customer name from the rest
        parts = quote_number.split(' ', 1)
        if len(parts) > 1:
            # Return the part after the first space (employee initials + date + letter)
            return parts[1]
        else:
            # If no space found, return the original quote number
            return quote_number
    
    def generate_unified_quote(
        self,
        quote_items: List[Dict[str, Any]],
        customer_name: str,
        attention_name: str,
        quote_number: str,
        output_path: str,
        employee_info: Optional[Dict[str, str]] = None,
        **kwargs
    ) -> bool:
        """
        Generate a unified quote with consistent formatting for single or multiple items.
        
        Args:
            quote_items: List of quote items
            customer_name: Customer company name
            attention_name: Contact person name
            quote_number: Quote number
            output_path: Output file path
            employee_info: Employee information dict with 'name', 'phone', 'email'
            **kwargs: Additional template variables
            
        Returns:
            True if successful, False otherwise
        """
        if not quote_items:
            logger.error("No quote items provided")
            return False
        
        try:
            # Get primary item for template selection
            primary_item = None
            for item in quote_items:
                if item.get('type') == 'main':
                    primary_item = item
                    break
            
            if not primary_item:
                logger.error("No main items found in quote")
                return False
            
            # Extract primary model
            primary_part_number = primary_item.get('part_number', '')
            primary_model = self.extract_model_from_part_number(primary_part_number)
            
            # Load base template
            template_path = self.get_base_template_path(primary_model)
            if not template_path:
                logger.error(f"No template found for primary model: {primary_model}")
                return False
            
            # Load the template document
            doc = Document(str(template_path))
            
            # Get employee information
            employee_name = ""
            employee_phone = ""
            employee_email = ""
            if employee_info:
                employee_name = employee_info.get('name', 'John Nicholosi')
                employee_phone = employee_info.get('phone', '(713) 467-4438')
                employee_email = employee_info.get('email', 'John@babbitt.us')
            else:
                employee_name = kwargs.get('employee_name', 'John Nicholosi')
                employee_phone = kwargs.get('employee_phone', '(713) 467-4438')
                employee_email = kwargs.get('employee_email', 'John@babbitt.us')
            
            # Prepare primary item variables (for template replacement)
            primary_data = primary_item.get('data', {})
            
            # Base template variables
            variables = {
                # Header information
                'date': datetime.now().strftime("%B %d, %Y"),
                'customer_name': customer_name,
                'company_name': customer_name,
                'attention_name': attention_name,
                'contact_name': attention_name,
                'quote_number': quote_number,
                # Create display quote number without customer name prefix
                'display_quote_number': self._extract_display_quote_number(quote_number),
                
                # Primary item information
                'part_number': primary_part_number,
                'quantity': str(primary_item.get('quantity', 1)),
                'unit_price': f"{primary_data.get('total_price', 0):.2f}",
                'price': f"{primary_data.get('total_price', 0):.2f}",
                'supply_voltage': primary_data.get('voltage', '115VAC'),
                'voltage': primary_data.get('voltage', '115VAC'),
                'probe_length': str(primary_data.get('probe_length', 12)),
                'length': str(primary_data.get('probe_length', 12)),
                
                # Technical specifications
                'process_connection_size': f"{primary_data.get('pc_size', '¾')}\"",
                'pc_type': primary_data.get('pc_type', 'NPT'),
                'pc_size': primary_data.get('pc_size', '¾"'),
                'pc_matt': primary_data.get('pc_matt', 'SS'),
                'insulator_material': self._extract_insulator_material_name(primary_data),
                'insulator_length': f"{primary_data.get('base_insulator_length', 4)}\"",
                'probe_material': primary_data.get('probe_material_name', '316SS'),
                'probe_diameter': primary_data.get('probe_diameter', '½"'),
                'max_temperature': f"{primary_data.get('max_temperature', 450)}°F",
                'max_pressure': f"{primary_data.get('max_pressure', 300)} PSI",
                'output_type': primary_data.get('output_type', '10 Amp SPDT Relay'),
                
                # Company information
                'company_contact': employee_name,
                'company_phone': employee_phone,
                'company_email': employee_email,
                'company_website': 'www.babbittinternational.com',
                'employee_name': employee_name,
                'employee_phone': employee_phone,
                'employee_email': employee_email,
                
                # Terms and conditions
                'delivery_terms': 'NET 30 W.A.C.',
                'fob_terms': 'FOB, Houston, TX',
                'quote_validity': '30 days',
                'lead_time': kwargs.get('lead_time', 'In Stock')
            }
            
            # Add any additional variables from kwargs
            variables.update(kwargs)
            
            # Replace template variables with primary item data
            self._replace_template_variables(doc, variables)
            
            # If multiple items, add additional sections
            if len(quote_items) > 1:
                self._add_multi_item_section(doc, quote_items, 1)
                self._add_quote_summary(doc, quote_items)
            
            # Save the document
            doc.save(output_path)
            logger.info(f"Unified quote generated successfully: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"Error generating unified quote: {e}")
            import traceback
            traceback.print_exc()
            return False

# Convenience function for easy integration
def generate_unified_quote(
    quote_items: List[Dict[str, Any]],
    customer_name: str,
    attention_name: str,
    quote_number: str,
    output_path: str,
    employee_info: Optional[Dict[str, str]] = None,
    **kwargs
) -> bool:
    """
    Generate a unified quote with consistent formatting.
    
    Args:
        quote_items: List of quote items
        customer_name: Customer company name
        attention_name: Contact person name
        quote_number: Quote number
        output_path: Output file path
        employee_info: Employee information dict with 'name', 'phone', 'email'
        **kwargs: Additional template variables
        
    Returns:
        True if successful, False otherwise
    """
    generator = UnifiedQuoteGenerator()
    return generator.generate_unified_quote(
        quote_items=quote_items,
        customer_name=customer_name,
        attention_name=attention_name,
        quote_number=quote_number,
        output_path=output_path,
        employee_info=employee_info,
        **kwargs
    )

# Test function
def test_unified_quote_generation():
    """Test the unified quote generation with sample data."""
    if not DOCX_AVAILABLE:
        print("❌ python-docx not available")
        return False
    
    # Test data for single item
    single_item_test = [{
        'type': 'main',
        'part_number': 'LS2000-115VAC-S-12',
        'quantity': 1,
        'data': {
            'model': 'LS2000',
            'voltage': '115VAC',
            'probe_length': 12,
            'probe_material_name': '316SS',
            'total_price': 1250.0,
            'pc_type': 'NPT',
            'pc_size': '¾',
            'max_temperature': 450,
            'max_pressure': 300,
            'output_type': '10 Amp SPDT Relay'
        }
    }]
    
    # Test data for multi-item
    multi_item_test = [
        {
            'type': 'main',
            'part_number': 'LS2000-115VAC-S-12',
            'quantity': 1,
            'data': {
                'model': 'LS2000',
                'voltage': '115VAC',
                'probe_length': 12,
                'probe_material_name': '316SS',
                'total_price': 1250.0,
                'pc_type': 'NPT',
                'pc_size': '¾',
                'max_temperature': 450,
                'max_pressure': 300,
                'output_type': '10 Amp SPDT Relay'
            }
        },
        {
            'type': 'main',
            'part_number': 'LS2100-24VDC-H-8',
            'quantity': 2,
            'data': {
                'model': 'LS2100',
                'voltage': '24VDC',
                'probe_length': 8,
                'probe_material_name': 'Halar',
                'total_price': 1350.0,
                'pc_type': 'NPT',
                'pc_size': '¾',
                'max_temperature': 450,
                'max_pressure': 300,
                'output_type': '10 Amp SPDT Relay'
            }
        },
        {
            'type': 'spare',
            'part_number': 'SPARE-PROBE-001',
            'quantity': 1,
            'data': {
                'description': 'Replacement Probe Assembly',
                'category': 'probe_assembly',
                'pricing': {
                    'total_price': 125.0
                }
            }
        }
    ]
    
    # Test single item
    print("Testing single item quote generation...")
    success1 = generate_unified_quote(
        quote_items=single_item_test,
        customer_name="ACME Industrial Solutions",
        attention_name="John Smith, P.E.",
        quote_number="Q-2024-TEST-001",
        output_path="test_unified_single.docx",
        employee_info={'name': 'John Nicholosi', 'phone': '(713) 467-4438', 'email': 'John@babbitt.us'}
    )
    
    # Test multi-item
    print("Testing multi-item quote generation...")
    success2 = generate_unified_quote(
        quote_items=multi_item_test,
        customer_name="ACME Industrial Solutions",
        attention_name="John Smith, P.E.",
        quote_number="Q-2024-TEST-002",
        output_path="test_unified_multi.docx",
        employee_info={'name': 'John Nicholosi', 'phone': '(713) 467-4438', 'email': 'John@babbitt.us'}
    )
    
    if success1 and success2:
        print("✅ Unified quote generation tests passed!")
        return True
    else:
        print("❌ Unified quote generation tests failed")
        return False

if __name__ == "__main__":
    test_unified_quote_generation()