"""
Master Unified Quote Generation System

This module provides a completely unified quote generation system that creates
consistent formatting for all quotes regardless of item count or model types.
Uses a programmatically built template structure instead of model-specific templates.
"""

import os
import re
from pathlib import Path
from typing import Dict, Any, Optional, List, TYPE_CHECKING
from datetime import datetime
import logging

if TYPE_CHECKING:
    from docx import Document

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.shared import OxmlElement, qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)

def _extract_display_quote_number(quote_number: str) -> str:
    """
    Extract display quote number without customer name prefix.
    
    Args:
        quote_number: Full quote number (e.g., "ACME ZF071925A")
        
    Returns:
        Display quote number without customer name (e.g., "ZF071925A")
    """
    # Split by space to separate customer name from the rest
    parts = quote_number.split(' ', 1)
    if len(parts) > 1:
        # Return the part after the first space (employee initials + date + letter)
        return parts[1]
    else:
        # If no space found, return the original quote number
        return quote_number

class MasterQuoteGenerator:
    """
    Master quote generator that creates completely unified quotes using
    a programmatically built template structure.
    """
    
    def __init__(self):
        """Initialize the master quote generator."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required. Install with: pip install python-docx")
    
    def _create_header_section(self, doc: Any, quote_data: Dict[str, Any]) -> None:
        """
        Create the standard header section with logo area, date, customer info, and quote number.
        
        Args:
            doc: Document object
            quote_data: Quote information dictionary
        """
        # Company logo/header area (placeholder for logo)
        header_para = doc.add_paragraph()
        header_run = header_para.add_run("BABBITT INTERNATIONAL")
        header_run.bold = True
        header_run.font.size = Pt(18)
        header_run.font.color.rgb = RGBColor(0, 0, 0)
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle_para = doc.add_paragraph()
        subtitle_run = subtitle_para.add_run("Point Level Switch Quote")
        subtitle_run.font.size = Pt(14)
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add some spacing
        doc.add_paragraph("")
        
        # Date, Customer, Quote Number section
        info_table = doc.add_table(rows=4, cols=2)
        info_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # Date
        date_cell = info_table.cell(0, 0)
        date_cell.text = "Date:"
        date_cell.paragraphs[0].runs[0].bold = True
        info_table.cell(0, 1).text = quote_data.get('date', datetime.now().strftime("%B %d, %Y"))
        
        # Customer
        customer_cell = info_table.cell(1, 0)
        customer_cell.text = "Customer:"
        customer_cell.paragraphs[0].runs[0].bold = True
        info_table.cell(1, 1).text = quote_data.get('customer_name', 'Customer Name')
        
        # Attention
        attn_cell = info_table.cell(2, 0)
        attn_cell.text = "Attention:"
        attn_cell.paragraphs[0].runs[0].bold = True
        info_table.cell(2, 1).text = quote_data.get('attention_name', 'Contact Person')
        
        # Quote Number
        quote_cell = info_table.cell(3, 0)
        quote_cell.text = "Quote Number:"
        quote_cell.paragraphs[0].runs[0].bold = True
        # Extract display quote number without customer name prefix
        full_quote_number = quote_data.get('quote_number', 'Q-XXXX-XXXX')
        display_quote_number = _extract_display_quote_number(full_quote_number)
        info_table.cell(3, 1).text = display_quote_number
        
        # Add spacing after header
        doc.add_paragraph("")
        doc.add_paragraph("")
    
    def _add_item_section(self, doc: Any, item: Dict[str, Any], item_number: int) -> None:
        """
        Add a consistent item section with all specifications.
        
        Args:
            doc: Document object
            item: Quote item data
            item_number: Item number (1, 2, 3, etc.)
        """
        # Item header
        item_header = doc.add_paragraph()
        header_text = f"ITEM {item_number}: {item.get('type', 'MAIN').upper()}"
        header_run = item_header.add_run(header_text)
        header_run.bold = True
        header_run.font.size = Pt(12)
        header_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Part Number and Quantity (prominent display)
        part_para = doc.add_paragraph()
        part_para.add_run("Part Number: ").bold = True
        part_para.add_run(item.get('part_number', 'N/A'))
        
        qty_para = doc.add_paragraph()
        qty_para.add_run("Quantity: ").bold = True
        qty_para.add_run(str(item.get('quantity', 1)))
        
        # Add detailed specifications based on item type
        if item.get('type') == 'main':
            self._add_main_item_specifications(doc, item)
        else:
            self._add_spare_part_specifications(doc, item)
        
        # Add spacing between items
        doc.add_paragraph("")
    
    def _add_main_item_specifications(self, doc: Any, item: Dict[str, Any]) -> None:
        """
        Add technical specifications for main items.
        
        Args:
            doc: Document object
            item: Main item data
        """
        item_data = item.get('data', {})
        
        # Technical specifications header
        spec_header = doc.add_paragraph()
        spec_run = spec_header.add_run("Technical Specifications:")
        spec_run.bold = True
        spec_run.font.size = Pt(11)
        
        # Create specifications list with consistent formatting
        specs = [
            ("Model", item_data.get('model', 'N/A')),
            ("Supply Voltage", item_data.get('voltage', 'N/A')),
            ("Probe Length", f"{item_data.get('probe_length', 'N/A')}\"" if item_data.get('probe_length') else 'N/A'),
            ("Probe Material", item_data.get('probe_material_name', item_data.get('probe_material', 'N/A'))),
            ("Probe Diameter", item_data.get('probe_diameter', 'N/A')),
            ("Insulator Material", self._get_insulator_material_name(item_data)),
            ("Insulator Length", f"{item_data.get('base_insulator_length', 4)}\""),
            ("Maximum Temperature", f"{item_data.get('max_temperature', 450)}°F"),
            ("Maximum Pressure", f"{item_data.get('max_pressure', 300)} PSI"),
            ("Process Connection", f"{item_data.get('pc_type', 'NPT')} {item_data.get('pc_size', '¾')}\""),
            ("Output Type", item_data.get('output_type', '10 Amp SPDT Relay'))
        ]
        
        # Add each specification with bullet points
        for spec_name, spec_value in specs:
            if spec_value and spec_value != 'N/A':
                spec_para = doc.add_paragraph()
                spec_para.add_run(f"• {spec_name}: ").bold = True
                spec_para.add_run(str(spec_value))
        
        # Add options if present
        options = item_data.get('options', [])
        if options:
            options_para = doc.add_paragraph()
            options_para.add_run("• Options: ").bold = True
            options_text = ', '.join([opt.split(':')[0] if ':' in opt else opt for opt in options])
            options_para.add_run(options_text)
        
        # Add pricing
        self._add_item_pricing(doc, item)
    
    def _add_spare_part_specifications(self, doc: Any, item: Dict[str, Any]) -> None:
        """
        Add specifications for spare parts.
        
        Args:
            doc: Document object
            item: Spare part item data
        """
        item_data = item.get('data', {})
        
        # Description
        desc_para = doc.add_paragraph()
        desc_para.add_run("• Description: ").bold = True
        desc_para.add_run(item_data.get('description', 'Spare Part'))
        
        # Category
        if item_data.get('category'):
            cat_para = doc.add_paragraph()
            cat_para.add_run("• Category: ").bold = True
            cat_para.add_run(item_data.get('category', 'General'))
        
        # Add pricing
        self._add_item_pricing(doc, item)
    
    def _add_item_pricing(self, doc: Any, item: Dict[str, Any]) -> None:
        """
        Add pricing information for an item.
        
        Args:
            doc: Document object
            item: Item data
        """
        # Get pricing based on item type
        if item.get('type') == 'main':
            unit_price = item.get('data', {}).get('total_price', 0)
        else:  # spare part
            unit_price = item.get('data', {}).get('pricing', {}).get('total_price', 0)
        
        quantity = item.get('quantity', 1)
        total_price = unit_price * quantity
        
        # Pricing section
        pricing_para = doc.add_paragraph()
        pricing_para.add_run("• Unit Price: ").bold = True
        pricing_para.add_run(f"${unit_price:.2f}")
        
        if quantity > 1:
            pricing_para.add_run(" | Total: ").bold = True
            pricing_para.add_run(f"${total_price:.2f}")
    
    def _add_quote_summary(self, doc: Any, quote_items: List[Dict[str, Any]]) -> None:
        """
        Add quote summary section with totals (only for multi-item quotes).
        
        Args:
            doc: Document object
            quote_items: List of all quote items
        """
        if len(quote_items) <= 1:
            return  # No summary needed for single items
        
        # Add some spacing
        doc.add_paragraph("")
        
        # Summary header
        summary_header = doc.add_paragraph()
        summary_run = summary_header.add_run("QUOTE SUMMARY")
        summary_run.bold = True
        summary_run.font.size = Pt(14)
        summary_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Create summary table
        table = doc.add_table(rows=1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header row
        header_cells = table.rows[0].cells
        headers = ['Item', 'Part Number', 'Type', 'Qty', 'Total Price']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            # Make header bold
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Add data rows
        total_quote_value = 0.0
        for i, item in enumerate(quote_items, 1):
            row_cells = table.add_row().cells
            
            # Item number
            row_cells[0].text = str(i)
            
            # Part number
            row_cells[1].text = item.get('part_number', 'N/A')
            
            # Type
            row_cells[2].text = item.get('type', 'main').upper()
            
            # Quantity
            row_cells[3].text = str(item.get('quantity', 1))
            
            # Price calculation
            if item.get('type') == 'main':
                unit_price = item.get('data', {}).get('total_price', 0)
            else:  # spare part
                unit_price = item.get('data', {}).get('pricing', {}).get('total_price', 0)
            
            quantity = item.get('quantity', 1)
            item_total = unit_price * quantity
            total_quote_value += item_total
            
            row_cells[4].text = f"${item_total:.2f}"
        
        # Add total row
        total_row = table.add_row().cells
        total_row[0].text = ""
        total_row[1].text = ""
        total_row[2].text = ""
        total_row[3].text = "TOTAL:"
        total_row[4].text = f"${total_quote_value:.2f}"
        
        # Make total row bold
        for cell in total_row:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
    
    def _add_footer_section(self, doc: Any, quote_data: Dict[str, Any]) -> None:
        """
        Add the standard footer section with delivery terms, thank you, and employee info.
        
        Args:
            doc: Document object
            quote_data: Quote information dictionary
        """
        # Add some spacing
        doc.add_paragraph("")
        doc.add_paragraph("")
        
        # Delivery and Terms section
        terms_header = doc.add_paragraph()
        terms_run = terms_header.add_run("DELIVERY & TERMS")
        terms_run.bold = True
        terms_run.font.size = Pt(12)
        terms_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Terms details
        terms_list = [
            ("Delivery:", quote_data.get('lead_time', 'Contact for Lead Time')),
            ("Terms:", quote_data.get('delivery_terms', 'NET 30 W.A.C.')),
            ("FOB:", quote_data.get('fob_terms', 'FOB, Houston, TX')),
            ("Quote Valid:", quote_data.get('quote_validity', '30 days'))
        ]
        
        for term_name, term_value in terms_list:
            term_para = doc.add_paragraph()
            term_para.add_run(f"{term_name} ").bold = True
            term_para.add_run(str(term_value))
        
        # Add spacing
        doc.add_paragraph("")
        
        # Thank you message
        thank_you = doc.add_paragraph()
        thank_you_run = thank_you.add_run("Thank you for your business!")
        thank_you_run.bold = True
        thank_you_run.font.size = Pt(12)
        thank_you.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Employee contact information
        employee_info = quote_data.get('employee_info', {})
        if employee_info and employee_info.get('name'):
            doc.add_paragraph("")
            
            contact_header = doc.add_paragraph()
            contact_run = contact_header.add_run("Contact Information:")
            contact_run.bold = True
            contact_run.font.size = Pt(11)
            
            # Employee details
            emp_name = employee_info.get('name', 'John Nicholosi')
            emp_phone = employee_info.get('phone', '(713) 467-4438')
            emp_email = employee_info.get('email', 'John@babbitt.us')
            
            name_para = doc.add_paragraph()
            name_para.add_run("Sales Representative: ").bold = True
            name_para.add_run(emp_name)
            
            phone_para = doc.add_paragraph()
            phone_para.add_run("Phone: ").bold = True
            phone_para.add_run(emp_phone)
            
            email_para = doc.add_paragraph()
            email_para.add_run("Email: ").bold = True
            email_para.add_run(emp_email)
        
        # Company website
        doc.add_paragraph("")
        website_para = doc.add_paragraph()
        website_run = website_para.add_run("www.babbittinternational.com")
        website_run.font.size = Pt(10)
        website_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _get_insulator_material_name(self, item_data: dict) -> str:
        """
        Extract insulator material name from item data.
        
        Args:
            item_data: Item data dictionary
            
        Returns:
            Insulator material name
        """
        # Check for custom insulator first
        insulator = item_data.get('insulator', {})
        if isinstance(insulator, dict) and insulator.get('material_name'):
            return insulator['material_name']
        
        # Check for insulator material field
        insulator_material = item_data.get('insulator_material')
        if insulator_material:
            # Map codes to display names
            material_codes = {
                'TEF': 'Teflon',
                'U': 'UHMWPE', 
                'UHMWPE': 'UHMWPE',
                'DEL': 'DELRIN',
                'PEEK': 'PEEK',
                'CER': 'Ceramic'
            }
            return material_codes.get(insulator_material, insulator_material)
        
        # Default fallback
        return 'UHMWPE'
    
    def generate_master_quote(
        self,
        quote_items: List[Dict[str, Any]],
        customer_name: str,
        attention_name: str,
        quote_number: str,
        output_path: str,
        employee_info: Optional[Dict[str, str]] = None,
        **kwargs
    ) -> bool:
        """
        Generate a completely unified quote using the master template approach.
        
        Args:
            quote_items: List of quote items
            customer_name: Customer company name
            attention_name: Contact person name
            quote_number: Quote number
            output_path: Output file path
            employee_info: Employee information dict with 'name', 'phone', 'email'
            **kwargs: Additional template variables
            
        Returns:
            True if successful, False otherwise
        """
        if not quote_items:
            logger.error("No quote items provided")
            return False
        
        try:
            # Create a new blank document
            doc = Document()
            
            # Prepare quote data
            quote_data = {
                'date': datetime.now().strftime("%B %d, %Y"),
                'customer_name': customer_name,
                'attention_name': attention_name,
                'quote_number': quote_number,
                'employee_info': employee_info or {},
                'lead_time': kwargs.get('lead_time', 'Contact for Lead Time'),
                'delivery_terms': kwargs.get('delivery_terms', 'NET 30 W.A.C.'),
                'fob_terms': kwargs.get('fob_terms', 'FOB, Houston, TX'),
                'quote_validity': kwargs.get('quote_validity', '30 days')
            }
            
            # Build the document structure
            self._create_header_section(doc, quote_data)
            
            # Add all items in order with consistent formatting
            for i, item in enumerate(quote_items, 1):
                self._add_item_section(doc, item, i)
            
            # Add summary for multi-item quotes
            self._add_quote_summary(doc, quote_items)
            
            # Add footer section
            self._add_footer_section(doc, quote_data)
            
            # Save the document
            doc.save(output_path)
            logger.info(f"Master quote generated successfully: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"Error generating master quote: {e}")
            import traceback
            traceback.print_exc()
            return False

# Convenience function for easy integration
def generate_master_quote(
    quote_items: List[Dict[str, Any]],
    customer_name: str,
    attention_name: str,
    quote_number: str,
    output_path: str,
    employee_info: Optional[Dict[str, str]] = None,
    **kwargs
) -> bool:
    """
    Generate a master unified quote with completely consistent formatting.
    
    Args:
        quote_items: List of quote items
        customer_name: Customer company name
        attention_name: Contact person name
        quote_number: Quote number
        output_path: Output file path
        employee_info: Employee information dict with 'name', 'phone', 'email'
        **kwargs: Additional template variables
        
    Returns:
        True if successful, False otherwise
    """
    generator = MasterQuoteGenerator()
    return generator.generate_master_quote(
        quote_items=quote_items,
        customer_name=customer_name,
        attention_name=attention_name,
        quote_number=quote_number,
        output_path=output_path,
        employee_info=employee_info,
        **kwargs
    )

# Test function
def test_master_quote_generation():
    """Test the master quote generation with various scenarios."""
    if not DOCX_AVAILABLE:
        print("[ERROR] python-docx not available")
        return False
    
    print("Testing Master Quote Generation...")
    
    # Test data for multi-item quote
    test_data = [
        {
            'type': 'main',
            'part_number': 'LS2000-115VAC-S-12',
            'quantity': 1,
            'data': {
                'model': 'LS2000',
                'voltage': '115VAC',
                'probe_length': 12,
                'probe_material_name': '316SS',
                'probe_diameter': '½"',
                'total_price': 1250.0,
                'pc_type': 'NPT',
                'pc_size': '¾',
                'max_temperature': 450,
                'max_pressure': 300,
                'output_type': '10 Amp SPDT Relay',
                'base_insulator_length': 4,
                'insulator_material': 'UHMWPE',
                'options': ['Standard Housing']
            }
        },
        {
            'type': 'main',
            'part_number': 'LS6000-230VAC-H-16',
            'quantity': 2,
            'data': {
                'model': 'LS6000',
                'voltage': '230VAC',
                'probe_length': 16,
                'probe_material_name': 'Halar',
                'probe_diameter': '½"',
                'total_price': 1450.0,
                'pc_type': 'NPT',
                'pc_size': '1',
                'max_temperature': 450,
                'max_pressure': 300,
                'output_type': '10 Amp SPDT Relay',
                'base_insulator_length': 6,
                'insulator_material': 'Teflon',
                'options': ['High Temperature', 'Halar Coating']
            }
        },
        {
            'type': 'spare',
            'part_number': 'SPARE-PROBE-001',
            'quantity': 1,
            'data': {
                'description': 'Replacement Probe Assembly',
                'category': 'probe_assembly',
                'pricing': {
                    'total_price': 125.0
                }
            }
        }
    ]
    
    success = generate_master_quote(
        quote_items=test_data,
        customer_name="ACME Industrial Solutions",
        attention_name="John Smith, P.E.",
        quote_number="Q-2024-MASTER-001",
        output_path="test_master_quote.docx",
        employee_info={'name': 'John Nicholosi', 'phone': '(713) 467-4438', 'email': 'John@babbitt.us'},
        lead_time='2 - 3 Weeks'
    )
    
    if success:
        print("[SUCCESS] Master quote generation test passed!")
        print("Generated: test_master_quote.docx")
        return True
    else:
        print("[FAILED] Master quote generation test failed!")
        return False

if __name__ == "__main__":
    test_master_quote_generation()