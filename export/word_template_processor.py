"""
Word Template Processor for Professional Quote Generation

This module processes Word templates (.docx) with template variables like {{customer_name}}
and maintains the original professional formatting and layout.
"""

import os
import re
from pathlib import Path
from typing import Dict, Any, Optional, TYPE_CHECKING, List
from datetime import datetime
import logging

if TYPE_CHECKING:
    from docx import Document

try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)

class WordTemplateProcessor:
    """
    Processes Word templates by replacing {{variable}} placeholders with actual values.
    Maintains original formatting and professional appearance.
    """
    
    def __init__(self, templates_dir: Optional[str] = None):
        """
        Initialize the Word template processor.
        
        Args:
            templates_dir: Path to the templates directory
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required. Install with: pip install python-docx")
            
        if templates_dir is None:
            templates_dir = os.path.join(os.path.dirname(__file__), 'templates')
        self.templates_dir = Path(templates_dir)
        
    def get_template_path(self, model: str) -> Optional[Path]:
        """
        Get the template file path for a given model.
        
        Args:
            model: Model name (e.g., 'LS2000H', 'FS10000S')
            
        Returns:
            Path to the template file, or None if not found
        """
        template_name = f"{model}_template.docx"
        template_path = self.templates_dir / template_name
        
        if template_path.exists():
            return template_path
        
        logger.warning(f"Template not found: {template_path}")
        return None
    
    def replace_variables_in_paragraph(self, paragraph, variables: Dict[str, str]) -> None:
        """
        Replace template variables in a paragraph while preserving formatting.
        
        Args:
            paragraph: docx paragraph object
            variables: Dictionary of variable name -> value mappings
        """
        # Get the full text of the paragraph
        full_text = paragraph.text
        
        # Find all template variables in the format {{variable_name}}
        variable_pattern = r'\{\{([^}]+)\}\}'
        matches = re.findall(variable_pattern, full_text)
        
        if not matches:
            return
        
        # Replace variables in the full text
        new_text = full_text
        for var_name in matches:
            placeholder = f"{{{{{var_name}}}}}"
            value = variables.get(var_name, f"{{{{MISSING: {var_name}}}}}")
            new_text = new_text.replace(placeholder, str(value))
        
        # Clear the paragraph and add the new text
        # This preserves the paragraph's formatting but not individual run formatting
        paragraph.clear()
        paragraph.add_run(new_text)

    def _process_conditional_content_document(self, doc, variables: Dict[str, str], model: Optional[str] = None) -> None:
        """
        Process conditional content at the document level to avoid breaking formatting.
        Handles {{if_option:option_code:content}} and {{if_multiple_items:content}} patterns.
        Args:
            doc: Document object
            variables: Dictionary of variables including option information
            model: The template/model name (e.g., 'FS10000'), or None
        """
        import re
        try:
            # Process option codes
            option_codes = variables.get('option_codes', [])
            print(f"DEBUG: Raw option_codes from variables: {option_codes}")
            logger.info(f"Raw option_codes from variables: {option_codes}")
            print(f"DEBUG: Type of option_codes: {type(option_codes)}")
            logger.info(f"Type of option_codes: {type(option_codes)}")
            
            if isinstance(option_codes, str):
                # Handle string representation of a list (like "['XSP']")
                if option_codes.startswith('[') and option_codes.endswith(']'):
                    # It's a string representation of a list, try to evaluate it safely
                    try:
                        import ast
                        option_codes = ast.literal_eval(option_codes)
                        print(f"DEBUG: Parsed string list to: {option_codes}")
                        logger.info(f"Parsed string list to: {option_codes}")
                    except:
                        # Fallback: split by comma
                        option_codes = [opt.strip() for opt in option_codes.split(',') if opt.strip()]
                        print(f"DEBUG: Fallback split by comma: {option_codes}")
                        logger.info(f"Fallback split by comma: {option_codes}")
                else:
                    # Regular comma-separated string
                    option_codes = [opt.strip() for opt in option_codes.split(',') if opt.strip()]
                    print(f"DEBUG: Converted string option_codes to list: {option_codes}")
                    logger.info(f"Converted string option_codes to list: {option_codes}")
            elif isinstance(option_codes, list):
                print(f"DEBUG: Option_codes is already a list: {option_codes}")
                logger.info(f"Option_codes is already a list: {option_codes}")
            else:
                print(f"DEBUG: Unexpected option_codes type: {type(option_codes)}, value: {option_codes}")
                logger.warning(f"Unexpected option_codes type: {type(option_codes)}, value: {option_codes}")
                option_codes = []
            
            print(f"DEBUG: Final option_codes for processing: {option_codes}")
            logger.info(f"Final option_codes for processing: {option_codes}")
            
            # Get multi-item information
            has_multiple_items = variables.get('has_multiple_items', 'false').lower() == 'true'
            item_count = int(variables.get('item_count', '1'))
            
            print(f"DEBUG: Multi-item info - has_multiple_items: {has_multiple_items}, item_count: {item_count}")
            logger.info(f"Multi-item info - has_multiple_items: {has_multiple_items}, item_count: {item_count}")
            
            # Process paragraphs for conditional content
            for paragraph in doc.paragraphs:
                try:
                    self._process_conditional_content_paragraph(paragraph, option_codes, model, has_multiple_items, item_count)
                except Exception as e:
                    logger.warning(f"Error processing conditional content in paragraph: {e}")
                    # Continue processing other paragraphs
            
            # Process tables for conditional content
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            try:
                                self._process_conditional_content_paragraph(paragraph, option_codes, model, has_multiple_items, item_count)
                            except Exception as e:
                                logger.warning(f"Error processing conditional content in table paragraph: {e}")
                                # Continue processing other paragraphs
            
            # Process headers and footers for conditional content
            for section in doc.sections:
                header = section.header
                for paragraph in header.paragraphs:
                    try:
                        self._process_conditional_content_paragraph(paragraph, option_codes, model, has_multiple_items, item_count)
                    except Exception as e:
                        logger.warning(f"Error processing conditional content in header: {e}")
                        # Continue processing other paragraphs
                
                footer = section.footer
                for paragraph in footer.paragraphs:
                    try:
                        self._process_conditional_content_paragraph(paragraph, option_codes, model, has_multiple_items, item_count)
                    except Exception as e:
                        logger.warning(f"Error processing conditional content in footer: {e}")
                        # Continue processing other paragraphs
                        
        except Exception as e:
            logger.error(f"Error in conditional content processing: {e}")
            # Don't let this break the entire template processing
            # Just log the error and continue

    def _process_conditional_content_paragraph(self, paragraph, option_codes: List[str], model: Optional[str] = None, 
                                             has_multiple_items: bool = False, item_count: int = 1) -> None:
        """
        Process conditional content in a single paragraph.
        Args:
            paragraph: Paragraph object
            option_codes: List of option codes present in the part number
            model: The template/model name (e.g., 'FS10000'), or None
            has_multiple_items: Whether this is a multi-item quote
            item_count: Number of items in the quote
        """
        import re
        try:
            full_text = paragraph.text
            
            # Process option-based conditionals: {{if_option:option_code:content}}
            option_pattern = r'\{\{if_option:([^:]+):([^}]+)\}\}'
            option_matches = list(re.finditer(option_pattern, full_text))
            
            # Process multi-item conditionals: {{if_multiple_items:content}}
            multi_item_pattern = r'\{\{if_multiple_items:([^}]+)\}\}'
            multi_item_matches = list(re.finditer(multi_item_pattern, full_text))
            
            # Process item-specific conditionals: {{if_item_2:content}}, {{if_item_3:content}}, etc.
            item_pattern = r'\{\{if_item_(\d+):([^}]+)\}\}'
            item_matches = list(re.finditer(item_pattern, full_text))
            
            # Combine all matches and sort by position (reverse order for processing)
            all_matches = []
            for match in option_matches:
                all_matches.append(('option', match))
            for match in multi_item_matches:
                all_matches.append(('multi_item', match))
            for match in item_matches:
                all_matches.append(('item', match))
            
            # Sort by position in reverse order to process from end to beginning
            all_matches.sort(key=lambda x: x[1].start(), reverse=True)
            
            if not all_matches:
                return
            
            # Process all conditional content blocks
            for match_type, match in all_matches:
                start_pos = match.start()
                end_pos = match.end()
                
                if match_type == 'option':
                    # Handle option-based conditionals
                    option_code = match.group(1)
                    content = match.group(2)
                    
                    if option_code in option_codes:
                        # Option is present, keep the content
                        replacement = content
                    else:
                        # Option is not present, remove the conditional block
                        replacement = ""
                
                elif match_type == 'multi_item':
                    # Handle multi-item conditionals
                    content = match.group(1)
                    
                    if has_multiple_items:
                        # Multiple items exist, keep the content
                        replacement = content
                    else:
                        # Single item, remove the conditional block
                        replacement = ""
                
                elif match_type == 'item':
                    # Handle item-specific conditionals
                    item_number = int(match.group(1))
                    content = match.group(2)
                    
                    if item_number <= item_count:
                        # Item exists, keep the content
                        replacement = content
                    else:
                        # Item doesn't exist, remove the conditional block
                        replacement = ""
                
                # Replace the conditional block with the appropriate content
                paragraph.clear()
                new_text = full_text[:start_pos] + replacement + full_text[end_pos:]
                paragraph.add_run(new_text)
                full_text = new_text
                
        except Exception as e:
            logger.error(f"Error processing conditional content in paragraph: {e}")
            # Don't let this break the paragraph processing
            # Just log the error and continue
    
    def replace_variables_in_table(self, table, variables: Dict[str, str]) -> None:
        """
        Replace template variables in a table while preserving formatting.
        
        Args:
            table: docx table object
            variables: Dictionary of variable name -> value mappings
        """
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    self.replace_variables_in_paragraph(paragraph, variables)
    
    def process_template(self, model: str, variables: Dict[str, Any]) -> Optional[Any]:
        """
        Process a Word template by replacing all template variables.
        
        Args:
            model: Model name
            variables: Dictionary of template variables
            
        Returns:
            Processed Document object, or None if template not found
        """
        template_path = self.get_template_path(model)
        if not template_path:
            logger.error(f"Template not found for model: {model}")
            return None
        
        try:
            logger.info(f"Processing template: {template_path}")
            logger.info(f"Template variables: {list(variables.keys())}")
            
            # Load the template document
            doc = Document(str(template_path))
            logger.info(f"Template loaded successfully, {len(doc.paragraphs)} paragraphs")
            
            # Convert all values to strings
            str_variables = {k: str(v) if v is not None else "" for k, v in variables.items()}
            
            # Process conditional content first (before regular variable replacement)
            print("DEBUG: Processing conditional content...")
            logger.info("Processing conditional content...")
            print(f"DEBUG: Template variables keys: {list(str_variables.keys())}")
            logger.info(f"Template variables keys: {list(str_variables.keys())}")
            if 'option_codes' in str_variables:
                print(f"DEBUG: Option codes found in variables: {str_variables['option_codes']}")
                logger.info(f"Option codes found in variables: {str_variables['option_codes']}")
            else:
                print("DEBUG: No option_codes found in variables!")
                logger.warning("No option_codes found in variables!")
            
            self._process_conditional_content_document(doc, str_variables, model)
            print("DEBUG: Conditional content processing completed")
            logger.info("Conditional content processing completed")
            
            # Process all paragraphs
            logger.info("Processing regular template variables...")
            for paragraph in doc.paragraphs:
                self.replace_variables_in_paragraph(paragraph, str_variables)
            
            # Process all tables
            for table in doc.tables:
                self.replace_variables_in_table(table, str_variables)
            
            # Process headers and footers
            for section in doc.sections:
                header = section.header
                for paragraph in header.paragraphs:
                    self.replace_variables_in_paragraph(paragraph, str_variables)
                
                footer = section.footer
                for paragraph in footer.paragraphs:
                    self.replace_variables_in_paragraph(paragraph, str_variables)
            
            logger.info("Template processing completed successfully")
            return doc
            
        except Exception as e:
            logger.error(f"Error processing template {template_path}: {e}")
            import traceback
            logger.error(f"Traceback: {traceback.format_exc()}")
            return None
    
    def save_document(self, doc: Any, output_path: str) -> bool:
        """
        Save a processed document to a file.
        
        Args:
            doc: Document object to save
            output_path: Output file path
            
        Returns:
            True if successful, False otherwise
        """
        try:
            doc.save(output_path)
            return True
        except Exception as e:
            logger.error(f"Error saving document to {output_path}: {e}")
            return False

def _format_fraction(size_str: str) -> str:
    """
    Convert regular fractions to Unicode fractions for better display.
    
    Args:
        size_str: Size string like "3/4"", "1/2"", "1""
        
    Returns:
        Formatted size string with Unicode fractions
    """
    if not size_str:
        return size_str
    
    # Remove quotes for processing
    clean_size = size_str.replace('"', '').replace("'", '')
    
    # Convert common fractions to Unicode
    fraction_map = {
        '1/2': '½',
        '1/4': '¼',
        '3/4': '¾',
        '1/8': '⅛',
        '3/8': '⅜',
        '5/8': '⅝',
        '7/8': '⅞',
        '1/3': '⅓',
        '2/3': '⅔',
        '1/5': '⅕',
        '2/5': '⅖',
        '3/5': '⅗',
        '4/5': '⅘',
        '1/6': '⅙',
        '5/6': '⅚'
    }
    
    # Check if it's a fraction we can convert
    if clean_size in fraction_map:
        return fraction_map[clean_size] + '"'
    
    # Return original with quotes restored
    return size_str

def _format_probe_length(probe_length: str) -> str:
    """
    Format probe length to not show decimals for whole numbers.
    
    Args:
        probe_length: Probe length as string (e.g., "10.0", "12", "22.5")
        
    Returns:
        Formatted probe length (e.g., "10", "12", "22.5")
    """
    try:
        length_float = float(probe_length)
        # If it's a whole number, return as integer string
        if length_float == int(length_float):
            return str(int(length_float))
        else:
            return str(length_float)
    except (ValueError, TypeError):
        return str(probe_length)

def _parse_insulator_for_template(kwargs: Dict[str, Any]) -> Dict[str, str]:
    """
    Parse insulator information into granular template variables.
    
    Args:
        kwargs: Dictionary containing insulator data
        
    Returns:
        Dictionary with parsed insulator variables for templates
    """
    import re
    
    # Check if there's a custom insulator specified in the part number
    insulator_display = kwargs.get('insulator', '')
    max_temp = kwargs.get('max_temperature', '450°F')
    
    # Parse material name and length from insulator display string
    # Format examples: "8.0\" Teflon", "4.0\" UHMWPE (Base: 4.0\")"
    material = 'UHMWPE'  # Default
    length_num = 4.0     # Default
    
    if insulator_display and '"' in insulator_display:
        # Extract length and material from strings like "8.0\" Teflon" or "4.0\" UHMWPE (Base: 4.0\")"
        
        # First, extract the length (number before the first quote)
        length_match = re.search(r'(\d+(?:\.\d+)?)"', insulator_display)
        if length_match:
            length_num = float(length_match.group(1))
        
        # Then extract the material name (text after the quote, before any parentheses)
        material_match = re.search(r'"\s*([^(]+?)(?:\s*\(|$)', insulator_display)
        if material_match:
            material = material_match.group(1).strip()
        else:
            # Fallback: try to extract material from the whole string
            parts = insulator_display.split('"')
            if len(parts) > 1:
                material_part = parts[1].strip()
                # Remove any base length information in parentheses
                if '(' in material_part:
                    material_part = material_part.split('(')[0].strip()
                material = material_part
    
    # If no custom insulator, use the default insulator_material and insulator_length
    if not insulator_display:
        insulator_material = kwargs.get('insulator_material', 'UHMWPE')
        insulator_length_raw = kwargs.get('insulator_length', '4"')
        
        # Parse material name (remove any extra formatting)
        material = insulator_material.strip()
        
        # Parse length (extract number from strings like "4\"" or "4 inches")
        length_match = re.search(r'(\d+(?:\.\d+)?)', str(insulator_length_raw))
        length_num = float(length_match.group(1)) if length_match else 4.0
    
    # Format length string
    length_str = str(int(length_num)) if length_num == int(length_num) else str(length_num)
    
    # Determine if "Long" should be included (4" or more)
    long_text = "Long" if length_num >= 4.0 else ""
    
    # Parse temperature rating (extract number from strings like "450°F" or "180 F")
    temp_match = re.search(r'(\d+)', str(max_temp))
    temp_rating = temp_match.group(1) if temp_match else "450"
    
    return {
        'material': material,
        'length': length_str + '"',
        'long_text': long_text,
        'temp_rating': temp_rating
    }

def generate_word_quote(
    model: str,
    customer_name: str,
    attention_name: str,
    quote_number: str,
    part_number: str,
    unit_price: str,
    supply_voltage: str,
    probe_length: str,
    output_path: str,
    **kwargs
) -> bool:
    """
    Generate a quote document using Word templates with template variables.
    
    Args:
        model: Model name
        customer_name: Customer company name
        attention_name: Contact person name
        quote_number: Quote number
        part_number: Full part number
        unit_price: Unit price
        supply_voltage: Supply voltage
        probe_length: Probe length in inches
        output_path: Output file path
        **kwargs: Additional template variables
        
    Returns:
        True if successful, False otherwise
    """
    if not DOCX_AVAILABLE:
        logger.error("python-docx not available")
        return False
    
    try:
        # Create template processor
        processor = WordTemplateProcessor()
        
        # Parse insulator information for granular template variables
        insulator_vars = _parse_insulator_for_template(kwargs)
        
        # Prepare template variables
        variables = {
            # Header information
            'date': datetime.now().strftime("%B %d, %Y"),
            'customer_name': customer_name,
            'company_name': customer_name,  # Alternative name
            'attention_name': attention_name,
            'contact_name': attention_name,  # Alternative name
            'quote_number': quote_number,
            
            # Product information
            'part_number': part_number,
            'quantity': kwargs.get('quantity', "1"),
            'unit_price': unit_price,
            'price': unit_price,  # Alternative name
            'supply_voltage': supply_voltage,
            'voltage': supply_voltage,  # Alternative name
            'probe_length': probe_length,
            'length': probe_length,  # Alternative name
            
            # Technical specifications
            'process_connection_size': kwargs.get('process_connection_size', '¾"'),
            'pc_type': kwargs.get('pc_type', 'NPT'),
            'pc_size': _format_fraction(kwargs.get('pc_size', '¾"')),
            'pc_matt': kwargs.get('pc_matt', 'SS'),
            'pc_rate': kwargs.get('pc_rate'),
            'insulator_material': kwargs.get('insulator_material', 'UHMPE'),
            'insulator_length': kwargs.get('insulator_length', '4"'),
            'probe_material': kwargs.get('probe_material', '316SS'),
            'probe_diameter': kwargs.get('probe_diameter', '½"'),
            'max_temperature': kwargs.get('max_temperature', '450°F'),
            'max_pressure': kwargs.get('max_pressure', '300 PSI'),
            
            # Length pricing information
            'length_adder': kwargs.get('length_adder', 0.0),
            'adder_per': kwargs.get('adder_per', 'none'),
            
            # Enhanced insulator template variables
            'ins_material': insulator_vars['material'],
            'ins_length': insulator_vars['length'], 
            'ins_long': insulator_vars['long_text'],
            'ins_temp': insulator_vars['temp_rating'],
            
            # Enhanced probe template variables  
            'probe_size': kwargs.get('probe_diameter', '½"').replace('"', '').replace("'", ''),
            'probe_material': kwargs.get('probe_material_name', kwargs.get('probe_material', '316SS')),
            'probe_length': _format_probe_length(probe_length),
            
            # Output specifications
            'output_type': kwargs.get('output_type', '10 Amp SPDT Relay'),
            
            # Company information (use employee info if provided, otherwise use defaults)
            'company_contact': kwargs.get('employee_name', 'John Nicholosi'),
            'company_phone': kwargs.get('employee_phone', '(713) 467-4438'),
            'company_email': kwargs.get('employee_email', 'John@babbitt.us'),
            'company_website': 'www.babbittinternational.com',
            
            # Employee information (for templates that use these specific variables)
            'employee_name': kwargs.get('employee_name', 'John Nicholosi'),
            'employee_phone': kwargs.get('employee_phone', '(713) 467-4438'),
            'employee_email': kwargs.get('employee_email', 'John@babbitt.us'),
            
            # Terms and conditions
            'delivery_terms': 'NET 30 W.A.C.',
            'fob_terms': 'FOB, Houston, TX',
            'quote_validity': '30 days',
            
            # Lead time
            'lead_time': kwargs.get('lead_time', 'In Stock')
        }
        
        # Add any additional variables from kwargs
        variables.update(kwargs)
        
        # Process the template
        doc = processor.process_template(model, variables)
        if not doc:
            logger.error(f"Failed to process template for model: {model}")
            return False
        
        # Save the document
        success = processor.save_document(doc, output_path)
        if success:
            logger.info(f"Quote generated successfully: {output_path}")
        
        return success
        
    except Exception as e:
        logger.error(f"Error generating quote: {e}")
        return False

def generate_multi_item_word_quote(
    quote_items: List[Dict[str, Any]],
    customer_name: str,
    attention_name: str,
    quote_number: str,
    output_path: str,
    employee_info: Optional[Dict[str, str]] = None,
    **kwargs
) -> bool:
    """
    Generate a multi-item quote document using Word templates with conditional variables.
    
    Args:
        quote_items: List of quote items with 'type', 'part_number', 'quantity', 'data'
        customer_name: Customer company name
        attention_name: Contact person name
        quote_number: Quote number
        output_path: Output file path
        employee_info: Employee information dict with 'name', 'phone', 'email'
        **kwargs: Additional template variables
        
    Returns:
        True if successful, False otherwise
    """
    if not DOCX_AVAILABLE:
        logger.error("python-docx not available")
        return False
    
    if not quote_items:
        logger.error("No quote items provided")
        return False
    
    try:
        # Create template processor
        processor = WordTemplateProcessor()
        
        # Determine primary model for template selection
        main_items = [item for item in quote_items if item.get('type') == 'main']
        if not main_items:
            logger.error("No main items found in quote")
            return False
        
        # Use the first main item's model for the base template
        primary_item = main_items[0]
        primary_part_number = primary_item.get('part_number', '')
        primary_model = primary_part_number.split('-')[0] if '-' in primary_part_number else primary_part_number[:6]
        
        # Calculate totals
        total_quote_value = 0.0
        item_count = len(quote_items)
        
        for item in quote_items:
            if item['type'] == 'main':
                unit_price = item['data'].get('total_price', 0.0)
            else:  # spare part
                unit_price = item['data'].get('pricing', {}).get('total_price', 0.0)
            quantity = item.get('quantity', 1)
            total_quote_value += unit_price * quantity
        
        # Prepare base template variables (from primary item)
        primary_data = primary_item.get('data', {})
        insulator_vars = _parse_insulator_for_template(primary_data)
        
        # Get employee information
        employee_name = ""
        employee_phone = ""
        employee_email = ""
        if employee_info:
            employee_name = employee_info.get('name', 'John Nicholosi')
            employee_phone = employee_info.get('phone', '(713) 467-4438')
            employee_email = employee_info.get('email', 'John@babbitt.us')
        else:
            employee_name = kwargs.get('employee_name', 'John Nicholosi')
            employee_phone = kwargs.get('employee_phone', '(713) 467-4438')
            employee_email = kwargs.get('employee_email', 'John@babbitt.us')
        
        # Base template variables (same as single item)
        variables = {
            # Header information
            'date': datetime.now().strftime("%B %d, %Y"),
            'customer_name': customer_name,
            'company_name': customer_name,
            'attention_name': attention_name,
            'contact_name': attention_name,
            'quote_number': quote_number,
            
            # Primary item information (first main item)
            'part_number': primary_part_number,
            'quantity': str(primary_item.get('quantity', 1)),
            'unit_price': f"{primary_data.get('total_price', 0):.2f}",
            'price': f"{primary_data.get('total_price', 0):.2f}",
            'supply_voltage': primary_data.get('voltage', '115VAC'),
            'voltage': primary_data.get('voltage', '115VAC'),
            'probe_length': str(primary_data.get('probe_length', 12)),
            'length': str(primary_data.get('probe_length', 12)),
            
            # Technical specifications (from primary item)
            'process_connection_size': f"{primary_data.get('pc_size', '¾')}\"",
            'pc_type': primary_data.get('pc_type', 'NPT'),
            'pc_size': _format_fraction(primary_data.get('pc_size', '¾"')),
            'pc_matt': primary_data.get('pc_matt', 'SS'),
            'pc_rate': primary_data.get('pc_rate'),
            'insulator_material': _extract_insulator_material_name(primary_data),
            'insulator_length': f"{primary_data.get('base_insulator_length', 4)}\"",
            'probe_material': primary_data.get('probe_material_name', '316SS'),
            'probe_diameter': primary_data.get('probe_diameter', '½"'),
            'max_temperature': f"{primary_data.get('max_temperature', 450)}°F",
            'max_pressure': f"{primary_data.get('max_pressure', 300)} PSI",
            
            # Length pricing information
            'length_adder': primary_data.get('length_adder', 0.0),
            'adder_per': primary_data.get('adder_per', 'none'),
            
            # Enhanced insulator template variables
            'ins_material': insulator_vars['material'],
            'ins_length': insulator_vars['length'],
            'ins_long': insulator_vars['long_text'],
            'ins_temp': insulator_vars['temp_rating'],
            
            # Enhanced probe template variables
            'probe_size': primary_data.get('probe_diameter', '½"').replace('"', '').replace("'", ''),
            'probe_material': primary_data.get('probe_material_name', '316SS'),
            'probe_length': _format_probe_length(str(primary_data.get('probe_length', 12))),
            
            # Output specifications
            'output_type': primary_data.get('output_type', '10 Amp SPDT Relay'),
            
            # Company information
            'company_contact': employee_name,
            'company_phone': employee_phone,
            'company_email': employee_email,
            'company_website': 'www.babbittinternational.com',
            
            # Employee information
            'employee_name': employee_name,
            'employee_phone': employee_phone,
            'employee_email': employee_email,
            
            # Terms and conditions
            'delivery_terms': 'NET 30 W.A.C.',
            'fob_terms': 'FOB, Houston, TX',
            'quote_validity': '30 days',
            
            # Lead time
            'lead_time': kwargs.get('lead_time', 'In Stock'),
            
            # Multi-item specific variables
            'item_count': str(item_count),
            'total_quote_value': f"{total_quote_value:.2f}",
            'has_multiple_items': str(item_count > 1).lower(),
            'is_single_item': str(item_count == 1).lower()
        }
        
        # Add conditional variables for additional items (items 2-10)
        for i in range(2, min(11, item_count + 1)):
            if i <= len(quote_items):
                item = quote_items[i - 1]  # 0-based index
                item_data = item.get('data', {})
                
                # Basic item info
                variables[f'item_{i}_part_number'] = item.get('part_number', '')
                variables[f'item_{i}_type'] = item.get('type', '').upper()
                variables[f'item_{i}_quantity'] = str(item.get('quantity', 1))
                
                if item['type'] == 'main':
                    # Main item specifications
                    variables[f'item_{i}_model'] = item_data.get('model', '')
                    variables[f'item_{i}_voltage'] = item_data.get('voltage', '')
                    variables[f'item_{i}_probe_length'] = str(item_data.get('probe_length', ''))
                    variables[f'item_{i}_probe_material'] = item_data.get('probe_material_name', '')
                    variables[f'item_{i}_insulator_material'] = _extract_insulator_material_name(item_data)
                    variables[f'item_{i}_insulator_length'] = f"{item_data.get('base_insulator_length', 4)}\""
                    variables[f'item_{i}_max_temperature'] = f"{item_data.get('max_temperature', 450)}°F"
                    variables[f'item_{i}_max_pressure'] = f"{item_data.get('max_pressure', 300)} PSI"
                    variables[f'item_{i}_pc_type'] = item_data.get('pc_type', 'NPT')
                    variables[f'item_{i}_pc_size'] = f"{item_data.get('pc_size', '¾')}\""
                    variables[f'item_{i}_output_type'] = item_data.get('output_type', '10 Amp SPDT Relay')
                    variables[f'item_{i}_unit_price'] = f"{item_data.get('total_price', 0):.2f}"
                    variables[f'item_{i}_total_price'] = f"{item_data.get('total_price', 0) * item.get('quantity', 1):.2f}"
                    
                    # Options
                    options = item_data.get('options', [])
                    options_text = ', '.join([opt.split(':')[0] if ':' in opt else opt for opt in options])
                    variables[f'item_{i}_options'] = options_text
                    
                else:
                    # Spare part specifications
                    variables[f'item_{i}_description'] = item_data.get('description', '')
                    variables[f'item_{i}_unit_price'] = f"{item_data.get('pricing', {}).get('total_price', 0):.2f}"
                    variables[f'item_{i}_total_price'] = f"{item_data.get('pricing', {}).get('total_price', 0) * item.get('quantity', 1):.2f}"
                    variables[f'item_{i}_category'] = item_data.get('category', '')
            else:
                # Clear variables for non-existent items
                variables[f'item_{i}_part_number'] = ''
                variables[f'item_{i}_type'] = ''
                variables[f'item_{i}_quantity'] = ''
                variables[f'item_{i}_unit_price'] = ''
                variables[f'item_{i}_total_price'] = ''
        
        # Add any additional variables from kwargs
        variables.update(kwargs)
        
        # Process the template
        doc = processor.process_template(primary_model, variables)
        if not doc:
            logger.error(f"Failed to process template for model: {primary_model}")
            return False
        
        # Save the document
        success = processor.save_document(doc, output_path)
        if success:
            logger.info(f"Multi-item quote generated successfully: {output_path}")
        
        return success
        
    except Exception as e:
        logger.error(f"Error generating multi-item quote: {e}")
        import traceback
        traceback.print_exc()
        return False


def _extract_insulator_material_name(quote_data: dict) -> str:
    """
    Extract insulator material name from quote data.
    
    Args:
        quote_data: Quote data dictionary
        
    Returns:
        Insulator material name
    """
    # Check for custom insulator first
    insulator = quote_data.get('insulator', {})
    if isinstance(insulator, dict) and insulator.get('material_name'):
        return insulator['material_name']
    
    # Check for insulator material field
    insulator_material = quote_data.get('insulator_material')
    if insulator_material:
        return insulator_material
    
    # Default fallback
    return 'UHMWPE'


def generate_composed_multi_item_quote(
    quote_items: List[Dict[str, Any]],
    customer_name: str,
    attention_name: str,
    quote_number: str,
    output_path: str,
    employee_info: Optional[Dict[str, str]] = None,
    **kwargs
) -> bool:
    """
    Generate a multi-item quote by composing individual model-specific sections.
    This approach maintains the professional formatting and model-specific content
    for each item while combining them into one cohesive document.
    
    Args:
        quote_items: List of quote items with 'type', 'part_number', 'quantity', 'data'
        customer_name: Customer company name
        attention_name: Contact person name
        quote_number: Quote number
        output_path: Output file path
        employee_info: Employee information dict with 'name', 'phone', 'email'
        **kwargs: Additional template variables
        
    Returns:
        True if successful, False otherwise
    """
    if not DOCX_AVAILABLE:
        logger.error("python-docx not available")
        return False
    
    if not quote_items:
        logger.error("No quote items provided")
        return False
    
    try:
        # Create template processor
        processor = WordTemplateProcessor()
        
        # Calculate totals
        total_quote_value = 0.0
        item_count = len(quote_items)
        
        for item in quote_items:
            if item['type'] == 'main':
                unit_price = item['data'].get('total_price', 0.0)
            else:  # spare part
                unit_price = item['data'].get('pricing', {}).get('total_price', 0.0)
            quantity = item.get('quantity', 1)
            total_quote_value += unit_price * quantity
        
        # Get employee information
        employee_name = ""
        employee_phone = ""
        employee_email = ""
        if employee_info:
            employee_name = employee_info.get('name', 'John Nicholosi')
            employee_phone = employee_info.get('phone', '(713) 467-4438')
            employee_email = employee_info.get('email', 'John@babbitt.us')
        else:
            employee_name = kwargs.get('employee_name', 'John Nicholosi')
            employee_phone = kwargs.get('employee_phone', '(713) 467-4438')
            employee_email = kwargs.get('employee_email', 'John@babbitt.us')
        
        # Create a new document for the composed quote
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.section import WD_SECTION_START
        
        composed_doc = Document()
        
        # Add company header
        header = composed_doc.add_heading('BABBITT INTERNATIONAL', 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = composed_doc.add_heading('Point Level Switch Quote', 1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add quote header information
        header_info = composed_doc.add_paragraph()
        header_info.add_run(f"Date: ").bold = True
        header_info.add_run(datetime.now().strftime("%B %d, %Y"))
        
        header_info.add_run(f"\nCustomer: ").bold = True
        header_info.add_run(customer_name)
        
        header_info.add_run(f"\nAttention: ").bold = True
        header_info.add_run(attention_name)
        
        header_info.add_run(f"\nQuote Number: ").bold = True
        header_info.add_run(quote_number)
        
        if item_count > 1:
            header_info.add_run(f"\nTotal Items: ").bold = True
            header_info.add_run(str(item_count))
        
        composed_doc.add_paragraph("=" * 80)
        
        # Process each item and add its content
        for i, item in enumerate(quote_items, 1):
            item_type = item.get('type', 'main')
            part_number = item.get('part_number', '')
            quantity = item.get('quantity', 1)
            
            # Add item header
            item_header = composed_doc.add_heading(f"Item {i}: {item_type.upper()} PART", 2)
            
            # Add basic item info
            item_info = composed_doc.add_paragraph()
            item_info.add_run(f"Part Number: ").bold = True
            item_info.add_run(part_number)
            
            item_info.add_run(f"\nQuantity: ").bold = True
            item_info.add_run(str(quantity))
            
            if item_type == 'main':
                # Process main item using its specific model template
                success = _add_main_item_section(composed_doc, item, i, processor, 
                                               employee_name, employee_phone, employee_email)
                if not success:
                    logger.warning(f"Failed to process main item {i}: {part_number}")
                    # Add basic info as fallback
                    _add_fallback_item_info(composed_doc, item, i)
                    
            else:
                # Add spare part information
                _add_spare_part_section(composed_doc, item, i)
            
            # Add separator between items (except for last item)
            if i < len(quote_items):
                composed_doc.add_paragraph("-" * 60)
        
        # Add summary section for multi-item quotes
        if item_count > 1:
            composed_doc.add_paragraph("=" * 80)
            summary_heading = composed_doc.add_heading("QUOTE SUMMARY", 2)
            summary_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Create summary table
            summary_table = composed_doc.add_table(rows=1, cols=5)
            summary_table.style = 'Table Grid'
            
            # Header row
            hdr_cells = summary_table.rows[0].cells
            hdr_cells[0].text = 'Item'
            hdr_cells[1].text = 'Part Number'
            hdr_cells[2].text = 'Type'
            hdr_cells[3].text = 'Quantity'
            hdr_cells[4].text = 'Total Price'
            
            # Add data rows
            for i, item in enumerate(quote_items, 1):
                row_cells = summary_table.add_row().cells
                row_cells[0].text = str(i)
                row_cells[1].text = item.get('part_number', '')
                row_cells[2].text = item.get('type', '').upper()
                row_cells[3].text = str(item.get('quantity', 1))
                
                if item['type'] == 'main':
                    unit_price = item['data'].get('total_price', 0.0)
                else:
                    unit_price = item['data'].get('pricing', {}).get('total_price', 0.0)
                quantity = item.get('quantity', 1)
                total_price = unit_price * quantity
                row_cells[4].text = f"${total_price:.2f}"
            
            # Add total
            total_row = summary_table.add_row().cells
            total_row[0].text = ""
            total_row[1].text = ""
            total_row[2].text = ""
            total_row[3].text = "TOTAL:"
            total_row[4].text = f"${total_quote_value:.2f}"
            
            composed_doc.add_paragraph("")
        
        # Add footer
        composed_doc.add_paragraph("=" * 80)
        footer = composed_doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run("Babbitt International\n").bold = True
        footer.add_run(f"{employee_name}\n")
        footer.add_run(f"Phone: {employee_phone}\n")
        footer.add_run(f"Email: {employee_email}\n")
        footer.add_run("www.babbittinternational.com")
        
        # Add terms
        terms = composed_doc.add_paragraph()
        terms.alignment = WD_ALIGN_PARAGRAPH.CENTER
        terms.add_run("Terms: NET 30 W.A.C. | FOB: Houston, TX | Quote Valid: 30 days")
        
        # Save the composed document
        composed_doc.save(output_path)
        logger.info(f"Composed multi-item quote generated successfully: {output_path}")
        
        return True
        
    except Exception as e:
        logger.error(f"Error generating composed multi-item quote: {e}")
        import traceback
        traceback.print_exc()
        return False


def _add_main_item_section(doc, item: Dict[str, Any], item_number: int, processor: WordTemplateProcessor,
                          employee_name: str, employee_phone: str, employee_email: str) -> bool:
    """
    Add a main item section to the document using its specific model template.
    
    Args:
        doc: Document to add content to
        item: Quote item data
        item_number: Item number (1, 2, 3, etc.)
        processor: Template processor
        employee_name: Employee name
        employee_phone: Employee phone
        employee_email: Employee email
        
    Returns:
        True if successful, False otherwise
    """
    try:
        part_number = item.get('part_number', '')
        item_data = item.get('data', {})
        quantity = item.get('quantity', 1)
        
        # Extract model from part number
        model = part_number.split('-')[0] if '-' in part_number else part_number[:6]
        
        # Check if template exists for this model
        template_path = processor.get_template_path(model)
        if not template_path:
            logger.warning(f"No template found for model: {model}")
            return False
        
        # Prepare variables for this specific item
        insulator_vars = _parse_insulator_for_template(item_data)
        
        variables = {
            # Header information
            'date': datetime.now().strftime("%B %d, %Y"),
            'customer_name': 'SEE MAIN QUOTE HEADER',
            'company_name': 'SEE MAIN QUOTE HEADER',
            'attention_name': 'SEE MAIN QUOTE HEADER',
            'contact_name': 'SEE MAIN QUOTE HEADER',
            'quote_number': 'SEE MAIN QUOTE HEADER',
            
            # Item-specific information
            'part_number': part_number,
            'quantity': str(quantity),
            'unit_price': f"{item_data.get('total_price', 0):.2f}",
            'price': f"{item_data.get('total_price', 0):.2f}",
            'supply_voltage': item_data.get('voltage', '115VAC'),
            'voltage': item_data.get('voltage', '115VAC'),
            'probe_length': str(item_data.get('probe_length', 12)),
            'length': str(item_data.get('probe_length', 12)),
            
            # Technical specifications
            'process_connection_size': f"{item_data.get('pc_size', '¾')}\"",
            'pc_type': item_data.get('pc_type', 'NPT'),
            'pc_size': _format_fraction(item_data.get('pc_size', '¾"')),
            'pc_matt': item_data.get('pc_matt', 'SS'),
            'pc_rate': item_data.get('pc_rate'),
            'insulator_material': _extract_insulator_material_name(item_data),
            'insulator_length': f"{item_data.get('base_insulator_length', 4)}\"",
            'probe_material': item_data.get('probe_material_name', '316SS'),
            'probe_diameter': item_data.get('probe_diameter', '½"'),
            'max_temperature': f"{item_data.get('max_temperature', 450)}°F",
            'max_pressure': f"{item_data.get('max_pressure', 300)} PSI",
            
            # Length pricing information
            'length_adder': item_data.get('length_adder', 0.0),
            'adder_per': item_data.get('adder_per', 'none'),
            
            # Enhanced insulator template variables
            'ins_material': insulator_vars['material'],
            'ins_length': insulator_vars['length'],
            'ins_long': insulator_vars['long_text'],
            'ins_temp': insulator_vars['temp_rating'],
            
            # Enhanced probe template variables
            'probe_size': item_data.get('probe_diameter', '½"').replace('"', '').replace("'", ''),
            'probe_material': item_data.get('probe_material_name', '316SS'),
            'probe_length': _format_probe_length(str(item_data.get('probe_length', 12))),
            
            # Output specifications
            'output_type': item_data.get('output_type', '10 Amp SPDT Relay'),
            
            # Company information
            'company_contact': employee_name,
            'company_phone': employee_phone,
            'company_email': employee_email,
            'company_website': 'www.babbittinternational.com',
            
            # Employee information
            'employee_name': employee_name,
            'employee_phone': employee_phone,
            'employee_email': employee_email,
            
            # Terms and conditions
            'delivery_terms': 'NET 30 W.A.C.',
            'fob_terms': 'FOB, Houston, TX',
            'quote_validity': '30 days',
            
            # Options
            'option_codes': [opt.split(':')[0] if ':' in opt else opt for opt in item_data.get('options', [])]
        }
        
        # Process the template for this specific item
        item_doc = processor.process_template(model, variables)
        if not item_doc:
            logger.error(f"Failed to process template for model: {model}")
            return False
        
        # Extract content from the processed template (skip header sections)
        # We'll add the technical specifications and product details
        content_added = False
        
        for paragraph in item_doc.paragraphs:
            text = paragraph.text.strip()
            
            # Skip header information (customer, date, quote number)
            if any(skip in text.lower() for skip in ['customer', 'date', 'quote number', 'attention']):
                continue
            
            # Skip empty paragraphs
            if not text:
                continue
            
            # Add technical specifications and product details
            if any(keyword in text.lower() for keyword in ['specification', 'technical', 'model', 'voltage', 'probe', 'insulator', 'temperature', 'pressure', 'connection', 'output']):
                # Copy the paragraph to our document
                new_para = doc.add_paragraph()
                for run in paragraph.runs:
                    new_run = new_para.add_run(run.text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                content_added = True
        
        # If no content was extracted, add basic specifications
        if not content_added:
            _add_basic_specifications(doc, item_data)
        
        return True
        
    except Exception as e:
        logger.error(f"Error adding main item section: {e}")
        return False


def _add_spare_part_section(doc, item: Dict[str, Any], item_number: int) -> None:
    """
    Add a spare part section to the document.
    
    Args:
        doc: Document to add content to
        item: Quote item data
        item_number: Item number
    """
    item_data = item.get('data', {})
    
    # Add spare part specifications
    specs = doc.add_paragraph()
    specs.add_run("Description: ").bold = True
    specs.add_run(item_data.get('description', 'Spare Part'))
    
    specs.add_run(f"\nCategory: ").bold = True
    specs.add_run(item_data.get('category', 'General'))
    
    specs.add_run(f"\nUnit Price: ").bold = True
    unit_price = item_data.get('pricing', {}).get('total_price', 0.0)
    specs.add_run(f"${unit_price:.2f}")
    
    specs.add_run(f"\nQuantity: ").bold = True
    specs.add_run(str(item.get('quantity', 1)))
    
    specs.add_run(f"\nTotal Price: ").bold = True
    total_price = unit_price * item.get('quantity', 1)
    specs.add_run(f"${total_price:.2f}")


def _add_fallback_item_info(doc, item: Dict[str, Any], item_number: int) -> None:
    """
    Add basic item information as fallback when template processing fails.
    
    Args:
        doc: Document to add content to
        item: Quote item data
        item_number: Item number
    """
    item_data = item.get('data', {})
    
    info = doc.add_paragraph()
    info.add_run("Model: ").bold = True
    info.add_run(item_data.get('model', 'N/A'))
    
    info.add_run(f"\nVoltage: ").bold = True
    info.add_run(item_data.get('voltage', 'N/A'))
    
    info.add_run(f"\nProbe Length: ").bold = True
    info.add_run(f"{item_data.get('probe_length', 'N/A')}\"")
    
    info.add_run(f"\nUnit Price: ").bold = True
    info.add_run(f"${item_data.get('total_price', 0):.2f}")


def _add_basic_specifications(doc, item_data: Dict[str, Any]) -> None:
    """
    Add basic technical specifications when template content extraction fails.
    
    Args:
        doc: Document to add content to
        item_data: Item data dictionary
    """
    specs = doc.add_paragraph()
    specs.add_run("Technical Specifications:\n").bold = True
    
    specs.add_run(f"• Model: {item_data.get('model', 'N/A')}\n")
    specs.add_run(f"• Supply Voltage: {item_data.get('voltage', 'N/A')}\n")
    specs.add_run(f"• Probe Length: {item_data.get('probe_length', 'N/A')}\"\n")
    specs.add_run(f"• Probe Material: {item_data.get('probe_material_name', 'N/A')}\n")
    specs.add_run(f"• Insulator Material: {_extract_insulator_material_name(item_data)}\n")
    specs.add_run(f"• Maximum Temperature: {item_data.get('max_temperature', 'N/A')}°F\n")
    specs.add_run(f"• Maximum Pressure: {item_data.get('max_pressure', 'N/A')} PSI\n")
    specs.add_run(f"• Process Connection: {item_data.get('pc_type', 'N/A')} {item_data.get('pc_size', 'N/A')}\"\n")
    specs.add_run(f"• Output Type: {item_data.get('output_type', 'N/A')}\n")


# Test function
def test_template_processing():
    """Test the Word template processing with sample data."""
    if not DOCX_AVAILABLE:
        print("❌ python-docx not available")
        return False
    
    test_success = generate_word_quote(
        model="LS2000S",
        customer_name="ACME Industrial Solutions",
        attention_name="John Smith, P.E.",
        quote_number="Q-2024-TEST-001",
        part_number="LS2000-115VAC-S-12",
        unit_price="1,250.00",
        supply_voltage="115VAC",
        probe_length="12",
        output_path="test_word_quote.docx",
        insulator_material="Teflon",
        probe_material="316SS",
        max_temperature="450°F"
    )
    
    if test_success:
        print("✓ Test quote generated successfully!")
        return True
    else:
        print("❌ Test quote generation failed")
        return False

if __name__ == "__main__":
    test_template_processing() 