"""
Quote Generator for Babbitt Quote Generator
Creates Word documents from parsed quote data
"""

import os
from typing import Dict, Any, Optional
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class QuoteGenerator:
    def __init__(self):
        """Initialize the quote generator"""
        self.template_path = None
        
    def generate_quote(self, quote_data: Dict[str, Any], output_path: str) -> bool:
        """
        Generate a quote document from parsed data
        
        Args:
            quote_data: Parsed part number data
            output_path: Where to save the document
            
        Returns:
            bool: True if successful, False otherwise
        """
        
        if not DOCX_AVAILABLE:
            print("❌ python-docx not available - cannot generate Word documents")
            return False
        
        try:
            # Create a new document
            doc = Document()
            
            # Add header
            self._add_header(doc, quote_data)
            
            # Add part information
            self._add_part_info(doc, quote_data)
            
            # Add specifications
            self._add_specifications(doc, quote_data)
            
            # Add options
            self._add_options(doc, quote_data)
            
            # Add warnings/errors
            self._add_warnings_errors(doc, quote_data)
            
            # Add footer
            self._add_footer(doc)
            
            # Save document
            doc.save(output_path)
            
            print(f"✓ Quote generated: {output_path}")
            return True
            
        except Exception as e:
            print(f"❌ Error generating quote: {e}")
            return False
    
    def _add_header(self, doc: Document, data: Dict[str, Any]):
        """Add document header"""
        # Company header
        header = doc.add_heading('Babbitt International', 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_heading('Point Level Switch Quote', 1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Quote info
        quote_info = doc.add_paragraph()
        quote_info.add_run(f"Part Number: ").bold = True
        quote_info.add_run(data.get('original_part_number', 'N/A'))
        
        quote_info.add_run(f"\nDate: ").bold = True
        quote_info.add_run(datetime.now().strftime("%B %d, %Y"))
        
        doc.add_paragraph("=" * 60)
    
    def _add_part_info(self, doc: Document, data: Dict[str, Any]):
        """Add basic part information"""
        doc.add_heading('Product Information', 2)
        
        # Create table for specs
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Specification'
        hdr_cells[1].text = 'Value'
        
        # Add rows
        specs = [
            ('Model', data.get('model', 'N/A')),
            ('Voltage', data.get('voltage', 'N/A')),
            ('Probe Material', data.get('probe_material_name', 'N/A')),
            ('Probe Length', f"{data.get('probe_length', 'N/A')}\""),
            ('Housing', data.get('housing_type', 'N/A')),
            ('Output', data.get('output_type', 'N/A'))
        ]
        
        for spec_name, spec_value in specs:
            row_cells = table.add_row().cells
            row_cells[0].text = spec_name
            row_cells[1].text = str(spec_value)
        
        doc.add_paragraph("")
    
    def _add_specifications(self, doc: Document, data: Dict[str, Any]):
        """Add detailed specifications"""
        doc.add_heading('Technical Specifications', 2)
        
        # Process Connection
        doc.add_heading('Process Connection', 3)
        conn_p = doc.add_paragraph()
        
        if data.get('process_connection'):
            conn = data['process_connection']
            conn_p.add_run(f"Type: ").bold = True
            conn_p.add_run(f"{conn.get('type', 'N/A')}\n")
            conn_p.add_run(f"Size: ").bold = True
            conn_p.add_run(f"{conn.get('size', 'N/A')}\n")
            if conn.get('rating'):
                conn_p.add_run(f"Rating: ").bold = True
                conn_p.add_run(f"{conn.get('rating', 'N/A')}\n")
        else:
            conn_p.add_run(f"Type: ").bold = True
            conn_p.add_run(f"{data.get('process_connection_type', 'N/A')}\n")
            conn_p.add_run(f"Size: ").bold = True
            conn_p.add_run(f"{data.get('process_connection_size', 'N/A')}\n")
        
        conn_p.add_run(f"Material: ").bold = True
        conn_p.add_run(f"{data.get('process_connection_material', 'N/A')}")
        
        # Insulator
        doc.add_heading('Insulator', 3)
        ins_p = doc.add_paragraph()
        
        if data.get('insulator'):
            ins = data['insulator']
            ins_p.add_run(f"Material: ").bold = True
            ins_p.add_run(f"{ins.get('material_name', 'N/A')}\n")
            ins_p.add_run(f"Length: ").bold = True
            ins_p.add_run(f"{ins.get('length', 'N/A')}\"")
            
            # Add base length if different from actual length
            base_length = data.get('base_insulator_length')
            if base_length and ins.get('length') != base_length:
                ins_p.add_run(f"\nBase Length: ").bold = True
                ins_p.add_run(f"{base_length:.1f}\"")
        else:
            ins_p.add_run(f"Material: ").bold = True
            ins_p.add_run(f"{data.get('insulator_material', 'N/A')}\n")
            ins_p.add_run(f"Length: ").bold = True
            actual_length = data.get('insulator_length', 'N/A')
            ins_p.add_run(f"{actual_length}\"")
            
            # Add base length information
            base_length = data.get('base_insulator_length')
            if base_length:
                ins_p.add_run(f"\nBase Length: ").bold = True
                ins_p.add_run(f"{base_length:.1f}\"")
                
                # Add explanation of base length calculation
                probe_length = data.get('probe_length', 0)
                if probe_length >= 8.0:
                    rule = "Probe ≥8\": Base = 4\""
                elif probe_length >= 5.0:
                    rule = "Probe 5-7\": Base = 2\""
                else:
                    rule = "Probe ≤4\": Base = 1\""
                ins_p.add_run(f" ({rule})")
        
        # Operating Limits
        doc.add_heading('Operating Limits', 3)
        limits_p = doc.add_paragraph()
        limits_p.add_run(f"Max Temperature: ").bold = True
        limits_p.add_run(f"{data.get('max_temperature', 'N/A')}°F\n")
        limits_p.add_run(f"Max Pressure: ").bold = True
        limits_p.add_run(f"{data.get('max_pressure', 'N/A')} PSI\n")
        limits_p.add_run(f"O-Ring Material: ").bold = True
        limits_p.add_run(f"{data.get('oring_material', 'N/A')}")
    
    def _add_options(self, doc: Document, data: Dict[str, Any]):
        """Add options section"""
        options = data.get('options', [])
        
        if options:
            doc.add_heading('Options', 2)
            
            for option in options:
                option_p = doc.add_paragraph(style='List Bullet')
                option_p.add_run(f"{option.get('code', 'N/A')}: ").bold = True
                option_p.add_run(option.get('name', 'N/A'))
    
    def _add_warnings_errors(self, doc: Document, data: Dict[str, Any]):
        """Add warnings and errors"""
        warnings = data.get('warnings', [])
        errors = data.get('errors', [])
        
        if warnings:
            doc.add_heading('Warnings', 2)
            for warning in warnings:
                warning_p = doc.add_paragraph(style='List Bullet')
                warning_p.add_run("⚠️ ").bold = True
                warning_p.add_run(warning)
        
        if errors:
            doc.add_heading('Configuration Errors', 2)
            for error in errors:
                error_p = doc.add_paragraph(style='List Bullet')
                error_p.add_run("❌ ").bold = True
                error_p.add_run(error)
    
    def _add_footer(self, doc: Document):
        """Add document footer"""
        doc.add_paragraph("")
        doc.add_paragraph("=" * 60)
        
        footer_p = doc.add_paragraph()
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_p.add_run("Generated by Babbitt Quote Generator v1.0\n").italic = True
        footer_p.add_run(f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}").italic = True
        
        doc.add_paragraph("")
        contact_p = doc.add_paragraph()
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_p.add_run("Babbitt International\n").bold = True
        contact_p.add_run("www.babbittinternational.com")

# Test the generator if run directly
if __name__ == "__main__":
    print("Testing Quote Generator...")
    
    if not DOCX_AVAILABLE:
        print("❌ python-docx not available - install with: pip install python-docx")
        exit(1)
    
    # Test data
    test_data = {
        'original_part_number': 'LS2000-115VAC-S-10"-XSP-VR-8"TEFINS',
        'model': 'LS2000',
        'voltage': '115VAC',
        'probe_material_name': '316 Stainless Steel',
        'probe_length': 10.0,
        'housing_type': 'Cast Aluminum, NEMA 7, C, D; NEMA 9, E, F, & G',
        'output_type': '10 Amp SPDT Relay',
        'process_connection_type': 'NPT',
        'process_connection_size': '3/4"',
        'process_connection_material': 'S',
        'insulator': {
            'material_name': 'Teflon',
            'length': 8.0
        },
        'max_temperature': 450,
        'max_pressure': 300,
        'oring_material': 'Viton',
        'options': [
            {'code': 'XSP', 'name': 'Extra Static Protection'},
            {'code': 'VR', 'name': 'Vibration Resistance'}
        ],
        'warnings': ['This is a test warning'],
        'errors': [],
        'base_insulator_length': 4.0 # Added for testing base length
    }
    
    generator = QuoteGenerator()
    success = generator.generate_quote(test_data, "test_quote.docx")
    
    if success:
        print("✓ Test quote generated successfully!")
    else:
        print("❌ Failed to generate test quote") 